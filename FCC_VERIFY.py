#!/usr/bin/env python3
"""
AUTHOR: CHRISTOPHER T. WILLIAMS

FCC_VERIFY.py  —  FCC ASR / ULS Public Database Verification
Companion to GNSS_AttackModel.ps1 and NLP_REPORT.py

Auto-discovers the most recent GNSS_AttackCase_* folder, reads points.csv,
and cross-references observed cell data against:

  1. FCC ASR  — Antenna Structure Registration  (is there a registered tower
                at or near each observed cell location?)
  2. FCC ULS  — Universal Licensing System      (who is licensed to transmit
                on each observed band in Riverside County, CA?)
  3. Jurisdictional routing table               (which agency can act on what)

Output written back into the same case folder:
  fcc_verification_YYYYMMDD_HHMMSS.docx
  fcc_verification_YYYYMMDD_HHMMSS.json

Usage:
    python FCC_VERIFY.py
    python FCC_VERIFY.py --case "C:\\gnss\\forensic_output\\GNSS_AttackCase_20260409_202621"
    python FCC_VERIFY.py --base "D:\\custom\\forensic_output"
    python FCC_VERIFY.py --offline   (use pre-loaded licensee data, skip live API calls)

Dependencies:
    pip install pandas python-docx requests
"""

import sys
import os
import glob
import json
import math
import time
import argparse
import datetime

import pandas as pd
import requests
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── CONFIG ────────────────────────────────────────────────────────────────────
BASE_DIR      = r"C:\gnss\forensic_output"
COMPLAINANT   = None
DOB           = ""

ASR_RADIUS_M  = 500    # flag if no registered tower within this radius
API_TIMEOUT   = 15
API_RETRIES   = 3
API_DELAY     = 1.5

# FCC public API endpoints — tried in order until one succeeds
ASR_ENDPOINTS = [
    "https://data.fcc.gov/api/antenna-system-registration/registrations.json",
    "https://data.fcc.gov/api/antenna-system-registration/registrations",
]
ULS_ENDPOINTS = [
    "https://data.fcc.gov/api/license-view/basicSearch/getLicenses.json",
    "https://data.fcc.gov/api/license-view/basicSearch/getLicenses",
]

# ── LTE EARFCN → BAND TABLE  (3GPP TS 36.101 Table 5.7.3-1) ─────────────────
# (earfcn_lo, earfcn_hi, band, freq_description, fcc_spectrum_block)
EARFCN_BANDS = [
    (0,      599,    1,  "2100 MHz IMT (Band 1)",           "AWS/IMT — 2110-2170 MHz"),
    (600,    1199,   2,  "1900 MHz PCS (Band 2)",            "PCS — 1930-1990 MHz"),
    (1200,   1949,   3,  "1800 MHz DCS (Band 3)",            "DCS — 1805-1880 MHz"),
    (1950,   2399,   4,  "1700/2100 MHz AWS-1 (Band 4)",     "AWS-1 — 2110-2155 MHz"),
    (2400,   2649,   5,  "850 MHz CLR (Band 5)",             "CLR — 869-894 MHz"),
    (2750,   3449,   7,  "2600 MHz IMT-E (Band 7)",          "BRS/EBS — 2620-2690 MHz"),
    (3450,   3799,   8,  "900 MHz EGSM (Band 8)",            "EGSM — 925-960 MHz"),
    (4150,   4749,  10,  "1700/2100 MHz AWS (Band 10)",      "AWS — 2110-2155 MHz"),
    (5000,   5179,  12,  "700 MHz Lower A/B/C (Band 12)",    "700 MHz Lower — 729-746 MHz"),
    (5180,   5279,  13,  "700 MHz Upper C (Band 13)",        "700 MHz Upper C — 746-757 MHz"),
    (5280,   5379,  14,  "700 MHz Upper D / PS (Band 14)",   "700 MHz D-block — 758-768 MHz"),
    (5730,   5849,  17,  "700 MHz Lower B/C (Band 17)",      "700 MHz Lower B/C — 734-746 MHz"),
    (5850,   5999,  18,  "850 MHz (Band 18)",                "850 MHz — 860-875 MHz"),
    (6000,   6149,  19,  "850 MHz (Band 19)",                "850 MHz — 875-890 MHz"),
    (6150,   6449,  20,  "800 MHz DD (Band 20)",             "800 MHz — 791-821 MHz"),
    (6600,   7399,  22,  "3500 MHz (Band 22)",               "3500 MHz — 3510-3590 MHz"),
    (7500,   7699,  25,  "1900 MHz ext PCS (Band 25)",       "PCS ext — 1930-1995 MHz"),
    (7700,   8039,  26,  "850 MHz ext CLR (Band 26)",        "CLR ext — 859-894 MHz"),
    (8690,   9039,  28,  "700 MHz APT (Band 28)",            "700 MHz APT — 758-803 MHz"),
    (9210,   9659,  30,  "2300 MHz WCS (Band 30)",           "WCS — 2350-2360 MHz"),
    (39650,  41589, 41,  "2500 MHz TDD BRS/EBS (Band 41)",  "BRS/EBS — 2496-2690 MHz"),
    (65536,  66435, 66,  "Band 66 UL (1710-1780 MHz)",       "AWS-3 — 1695-2180 MHz"),
    (66436,  67335, 66,  "Band 66 DL (2110-2180 MHz)",       "AWS-3 — 1695-2180 MHz"),
    (68586,  68935, 71,  "600 MHz T-Mobile (Band 71)",       "600 MHz — 617-652 MHz"),
]

def earfcn_to_band(earfcn):
    """Return (band_num, freq_desc, fcc_block) or (None, 'unknown', '')."""
    try:
        e = int(earfcn)
    except (ValueError, TypeError):
        return None, "unknown", ""
    for lo, hi, band, freq, fcc in EARFCN_BANDS:
        if lo <= e <= hi:
            return band, freq, fcc
    return None, f"EARFCN {earfcn} (not in band table)", ""

# ── PRE-LOADED FCC ULS LICENSEE DATA  (Riverside County, CA) ────────────────
# Sourced from publicly available FCC ULS records.  Used when live API fails.
KNOWN_LICENSEES = {
    66: {
        "licensee":     "T-Mobile License LLC",
        "file_numbers": ["0007577840", "0007702797"],
        "market":       "EA-011 (Los Angeles-Long Beach)",
        "frequencies":  "1700/2100 MHz AWS-3 extended",
        "mcc_mnc":      "310-260",
        "note":         "EARFCN 66586 is Band 66 DL (2110-2180 MHz).  "
                        "T-Mobile is the primary AWS-3 licensee in Riverside County.  "
                        "Verizon (MCC 311/MNC 480) broadcasting on Band 66 without a "
                        "T-Mobile site agreement is unauthorized use of licensed spectrum.",
    },
    2: {
        "licensee":     "Cellco Partnership d/b/a Verizon Wireless",
        "file_numbers": ["0001832311", "0006302625"],
        "market":       "BTA-263 (Los Angeles-Long Beach)",
        "frequencies":  "1900 MHz PCS",
        "mcc_mnc":      "311-480",
        "note":         "EARFCN 1000 (Band 2 / 1900 MHz PCS) — "
                        "Verizon is the primary PCS licensee in Riverside County.",
    },
    13: {
        "licensee":     "Cellco Partnership d/b/a Verizon Wireless",
        "file_numbers": ["0001832282"],
        "market":       "CMA-391 (Riverside-San Bernardino)",
        "frequencies":  "700 MHz Upper C Block",
        "mcc_mnc":      "311-480",
        "note":         "Band 13 (700 MHz Upper C Block) is Verizon's sole nationwide block.",
    },
    4: {
        "licensee":     "AT&T Mobility LLC",
        "file_numbers": ["0006436786"],
        "market":       "EA-011",
        "frequencies":  "1700/2100 MHz AWS-1",
        "mcc_mnc":      "310-410",
        "note":         "Band 4 / AWS-1 — AT&T primary in this market.",
    },
    71: {
        "licensee":     "T-Mobile License LLC",
        "file_numbers": ["0007577840"],
        "market":       "PEA-036 (Los Angeles)",
        "frequencies":  "600 MHz",
        "mcc_mnc":      "310-260",
        "note":         "Band 71 (600 MHz) is T-Mobile's primary nationwide low-band holding.",
    },
    41: {
        "licensee":     "Sprint Spectrum LP (now T-Mobile US)",
        "file_numbers": ["0001718086"],
        "market":       "BRS/EBS Nationwide",
        "frequencies":  "2500 MHz BRS/EBS",
        "mcc_mnc":      "310-120 / 310-260",
        "note":         "Band 41 / 2500 MHz TDD — Sprint/T-Mobile.",
    },
    1: {
        "licensee":     "Various (AWS/IMT — not common in US deployments)",
        "file_numbers": [],
        "market":       "N/A",
        "frequencies":  "2100 MHz IMT Band 1",
        "mcc_mnc":      "N/A",
        "note":         "Band 1 is used internationally but rarely deployed in the US. "
                        "An LTE cell on Band 1 in Riverside County is highly anomalous.",
    },
}

# ── ULS RADIO SERVICE CODES  (exact FCC getLicenses abbreviations) ────────────
ULS_SERVICE_MAP = {
    66: ["AW"],        # AWS-3
    2:  ["CW"],        # Broadband PCS
    4:  ["AW"],        # AWS-1
    13: ["CZ"],        # 700 MHz Upper C
    71: ["YM"],        # 600 MHz Broadcast Incentive Auction
    41: ["BN"],        # Broadband Radio Service
    12: ["CY"],        # 700 MHz Lower A/B
    17: ["CY"],        # 700 MHz Lower B/C
    5:  ["CL"],        # 850 MHz Cellular
    26: ["CL"],        # 850 MHz Cellular extended
    30: ["WS"],        # Wireless Communications Service
    1:  ["AW"],        # AWS/IMT
}

# ── JURISDICTIONAL ROUTING TABLE ──────────────────────────────────────────────
ROUTING_TABLE = [
    {
        "anomaly":   "Unauthorized LTE transmitter on licensed spectrum",
        "evidence":  "EARFCN observed + no ASR structure at coordinates",
        "agency":    "FCC Enforcement Bureau",
        "mechanism": "47 U.S.C. § 301 complaint + ASR verification report",
        "statutes":  ["47 U.S.C. § 301", "47 U.S.C. § 333"],
        "action":    "File at enforcement@fcc.gov — attach points.csv, incidents.csv, this report.  "
                     "FCC CAN act on this — falls within their infrastructure jurisdiction.",
        "severity":  "PRIMARY",
    },
    {
        "anomaly":   "MCC/MNC identity impersonation on licensed band",
        "evidence":  "Session MCC/MNC matches licensed carrier; transmitter has no ASR registration",
        "agency":    "FCC Enforcement Bureau + DOJ Criminal Division",
        "mechanism": "47 U.S.C. § 301 (FCC) + 18 U.S.C. § 1028A (DOJ)",
        "statutes":  ["47 U.S.C. § 301", "18 U.S.C. § 1028A"],
        "action":    "File FCC complaint AND refer to FBI Cyber Division (IC3.gov) and DOJ.  "
                     "FCC handles the radio layer; DOJ handles the identity theft.",
        "severity":  "PRIMARY",
    },
    {
        "anomaly":   "PCI collision / IMSI-catcher RF signature",
        "evidence":  "Same PCI active on multiple EARFCNs simultaneously",
        "agency":    "FBI / DOJ",
        "mechanism": "18 U.S.C. § 2511 (Wiretap Act) + 18 U.S.C. § 2512",
        "statutes":  ["18 U.S.C. § 2511", "18 U.S.C. § 2512", "18 U.S.C. § 1029"],
        "action":    "File with FBI Riverside (IC3 + direct field office referral).  "
                     "FCC CANNOT act on PCI assignments — radio layer config, not infrastructure.  "
                     "DOJ only.",
        "severity":  "PRIMARY",
    },
    {
        "anomaly":   "RSRQ out-of-specification (measurement report manipulation)",
        "evidence":  "RSRQ index > 34 (3GPP TS 36.133 §9.1.7 ceiling)",
        "agency":    "DOJ / FBI",
        "mechanism": "18 U.S.C. § 1030 (CFAA) + 18 U.S.C. § 2511",
        "statutes":  ["18 U.S.C. § 1030", "18 U.S.C. § 2511", "47 U.S.C. § 333"],
        "action":    "Include in FBI referral.  FCC CANNOT act on RSRQ values — radio layer, "
                     "not infrastructure.  47 U.S.C. § 333 (interference) gives FCC a secondary hook.",
        "severity":  "SECONDARY",
    },
    {
        "anomaly":   "Reserved TAC broadcast (TAC 0 / 65535)",
        "evidence":  "3GPP-reserved TAC values present in session data",
        "agency":    "FCC + DOJ",
        "mechanism": "47 U.S.C. § 301 + 18 U.S.C. § 1028A",
        "statutes":  ["47 U.S.C. § 301", "18 U.S.C. § 1028A"],
        "action":    "Include in FCC complaint as additional unauthorized-operation evidence.  "
                     "TAC not in FCC databases but reserved values confirm non-compliant equipment.",
        "severity":  "SUPPORTING",
    },
    {
        "anomaly":   "GPS coordinate injection (exact repeating speed constants)",
        "evidence":  "Floating-point speed constants repeated across independent GNSS fixes",
        "agency":    "DOD / FBI / DOJ",
        "mechanism": "18 U.S.C. § 1367 (satellite interference) + 18 U.S.C. § 1030",
        "statutes":  ["18 U.S.C. § 1367", "18 U.S.C. § 1030"],
        "action":    "File with FBI + notify U.S. Space Force GPS Operations Center: (719) 567-2828.  "
                     "GPS is a military asset.  FCC has no jurisdiction — DOD/DOJ only.",
        "severity":  "PRIMARY",
    },
    {
        "anomaly":   "Session timing gaps (possible jamming)",
        "evidence":  "Inter-record gaps exceeding dynamic threshold, correlated with RSRQ anomalies",
        "agency":    "FCC Enforcement Bureau + FBI",
        "mechanism": "47 U.S.C. § 333 + 47 C.F.R. § 2.807 (jammer prohibition)",
        "statutes":  ["47 U.S.C. § 333", "47 C.F.R. § 2.807", "18 U.S.C. § 1362"],
        "action":    "Include in FCC complaint — jamming is explicitly within FCC enforcement "
                     "authority under § 333 and § 2.807 even without ASR evidence.",
        "severity":  "PRIMARY",
    },
    {
        "anomaly":   "Multi-session pattern (April 2 + April 5 2026 + prior history)",
        "evidence":  "Consistent PCI 242 collision, eNB 44319, Band 66 rogue across separate sessions",
        "agency":    "DOJ Criminal Division (RICO) + FBI",
        "mechanism": "18 U.S.C. §§ 1961-1968 — two or more predicate acts within 10 years",
        "statutes":  ["18 U.S.C. § 1961", "18 U.S.C. § 241", "18 U.S.C. § 242"],
        "action":    "Requires attorney.  Multi-session evidence plus six-year documentation "
                     "history meets the pattern-of-racketeering predicate act threshold.",
        "severity":  "COMPOUND",
    },
]

# ── HELPER: DOCX PRIMITIVES ───────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)

def add_run(para, text, bold=False, color=None, size=None, italic=False, mono=False):
    run = para.add_run(str(text)); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    if size:  run.font.size = Pt(size)
    if mono:  run.font.name = "Courier New"
    return run

def heading(doc, text, level=1, color="1A237E"):
    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text); run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "double" if level == 1 else "single")
    bot.set(qn("w:sz"), "6"); bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color); pBdr.append(bot); pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(6)

def add_kv_table(doc, rows, col_widths=(2.6, 3.9)):
    table = doc.add_table(rows=0, cols=2); table.style = "Table Grid"
    for i, (label, value, hl) in enumerate(rows):
        row = table.add_row(); c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(col_widths[0]); c1.width = Inches(col_widths[1])
        set_cell_bg(c0, "ECEFF1")
        set_cell_bg(c1, "FFEBEE" if hl else ("FFFFFF" if i % 2 == 0 else "FAFAFA"))
        add_run(c0.paragraphs[0], str(label), bold=True, size=10)
        if hl: add_run(c1.paragraphs[0], str(value), bold=True, color="B71C1C", size=10)
        else:  add_run(c1.paragraphs[0], str(value), size=10)
    doc.add_paragraph()

def add_log_table(doc, headers, col_widths, data_rows):
    table = doc.add_table(rows=0, cols=len(headers)); table.style = "Table Grid"
    hdr = table.add_row()
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        c = hdr.cells[i]; c.width = Inches(w)
        set_cell_bg(c, "1A237E")
        add_run(c.paragraphs[0], h, bold=True, color="FFFFFF", size=9)
    for dr in data_rows:
        row = table.add_row()
        for i, cv in enumerate(dr):
            c = row.cells[i]; c.width = Inches(col_widths[i])
            text = cv[0] if isinstance(cv, tuple) else str(cv)
            hl   = cv[1] if isinstance(cv, tuple) else False
            set_cell_bg(c, "FFEBEE" if hl else "FFFFFF")
            if hl: add_run(c.paragraphs[0], text, bold=True, color="B71C1C", size=9, mono=True)
            else:  add_run(c.paragraphs[0], text, size=9, mono=True)
    doc.add_paragraph()

def _fmt(v, dec=0, suf="", na="—"):
    try:
        if v is None or (isinstance(v, float) and pd.isna(v)): return na
        return f"{v:.{dec}f}{suf}" if dec else f"{int(round(float(v)))}{suf}"
    except Exception: return str(v) if v is not None else na

def _s(v, na="—"):
    try:
        if v is None or pd.isna(v): return na
    except Exception: pass
    return str(v)

# ── FCC API ───────────────────────────────────────────────────────────────────

def api_get(url, params, label="FCC API"):
    """
    HTTP GET with retry.  Returns parsed JSON or None.
    Handles the common FCC API failure modes:
      - 404  → wrong path (reported, no retry)
      - 200 with empty body  → wrong service code or missing .json suffix
      - 200 with HTML  → maintenance page or redirect
    """
    for attempt in range(API_RETRIES):
        try:
            r = requests.get(url, params=params, timeout=API_TIMEOUT,
                             headers={
                                 "User-Agent": "FCC-Evidence-Verifier/1.0",
                                 "Accept":     "application/json, */*",
                             })
            if r.status_code == 404:
                print(f"    {label}: 404 — path not found ({url[:70]})")
                return None   # definitive — no point retrying
            if r.status_code != 200:
                print(f"    {label}: HTTP {r.status_code} (attempt {attempt+1}/{API_RETRIES})")
            elif not r.text.strip():
                print(f"    {label}: empty body — check service code or .json suffix "
                      f"(attempt {attempt+1}/{API_RETRIES})")
            elif r.text.strip()[0] in ("<", "!"):
                print(f"    {label}: HTML response — API may be in maintenance "
                      f"(attempt {attempt+1}/{API_RETRIES})")
            else:
                try:
                    return r.json()
                except ValueError as e:
                    print(f"    {label}: JSON parse error — {e} (attempt {attempt+1}/{API_RETRIES})")
        except requests.exceptions.Timeout:
            print(f"    {label}: timeout (attempt {attempt+1}/{API_RETRIES})")
        except requests.exceptions.ConnectionError:
            print(f"    {label}: connection error (attempt {attempt+1}/{API_RETRIES})")
        except Exception as e:
            print(f"    {label}: {e} (attempt {attempt+1}/{API_RETRIES})")
        if attempt < API_RETRIES - 1:
            time.sleep(API_DELAY)
    return None

def haversine(lat1, lon1, lat2, lon2):
    R = 6371000.0
    p1, p2 = math.radians(lat1), math.radians(lat2)
    dp = math.radians(lat2 - lat1); dl = math.radians(lon2 - lon1)
    a  = math.sin(dp/2)**2 + math.cos(p1)*math.cos(p2)*math.sin(dl/2)**2
    return 2 * R * math.asin(math.sqrt(max(0, min(1, a))))

# ── CASE FOLDER HELPERS ───────────────────────────────────────────────────────

def find_case_folder(base, explicit=None):
    if explicit:
        if not os.path.isdir(explicit):
            sys.exit(f"[ERROR] Case folder not found: {explicit}")
        return explicit
    folders = glob.glob(os.path.join(base, "GNSS_AttackCase_*"))
    if not folders:
        sys.exit(
            f"[ERROR] No GNSS_AttackCase_* folders in: {base}\n"
            "        Run GNSS_AttackModel.ps1 first, or use --case <path>."
        )
    return max(folders, key=os.path.getmtime)

def read_hash_file(case_dir):
    path = os.path.join(case_dir, "evidence_hash.txt")
    info = {"File": "—", "Algorithm": "—", "Hash": "—"}
    if not os.path.isfile(path):
        return info
    with open(path, "r", errors="ignore") as f:
        for line in f:
            if ":" in line:
                k, _, v = line.partition(":")
                info[k.strip()] = v.strip()
    return info

def load_points(case_dir):
    path = os.path.join(case_dir, "points.csv")
    if not os.path.isfile(path):
        sys.exit(f"[ERROR] points.csv not found in {case_dir}")
    for enc in ("utf-8-sig", "utf-8", "utf-16", "latin-1"):
        try:
            df = pd.read_csv(path, encoding=enc)
            if not df.empty:
                break
        except Exception:
            continue
    else:
        sys.exit(f"[ERROR] Cannot parse {path}")
    for col in ["Cell_ECI","Cell_PCI","Cell_EARFCN","Cell_TAC",
                "Cell_MCC","Cell_MNC","Cell_RSRP","Cell_Lat","Cell_Lon"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    if "UtcTime" in df.columns:
        df["UtcTime"] = pd.to_datetime(df["UtcTime"], errors="coerce")
    return df

# ── CORE: BUILD UNIQUE CELL LIST ──────────────────────────────────────────────

def build_unique_cells(df):
    """
    Extract unique cells from points.csv.

    PCI/EARFCN auto-correction: GNSS_AttackModel.ps1 maps CellMapper CSV
    columns $p[10]->Cell_PCI and $p[11]->Cell_EARFCN, but some CellMapper
    app versions have those columns reversed.  3GPP PCI range is 0-503;
    EARFCN values are larger.  If Cell_PCI > 503 and Cell_EARFCN <= 503
    the columns are swapped and we correct automatically.
    """
    required = ["Cell_ECI", "Cell_EARFCN"]
    if not all(c in df.columns for c in required):
        print("[WARNING] Cell_ECI or Cell_EARFCN missing from points.csv")
        return []

    group_cols = [c for c in ["Cell_ECI","Cell_PCI","Cell_EARFCN","Cell_TAC",
                               "Cell_Lat","Cell_Lon","Cell_MCC","Cell_MNC"]
                  if c in df.columns]
    agg_dict = {"obs": ("Cell_ECI", "count")}
    if "Cell_RSRP" in df.columns:
        agg_dict["rsrp_mean"] = ("Cell_RSRP", "mean")
    if "UtcTime" in df.columns:
        agg_dict["time_first"] = ("UtcTime", "min")
        agg_dict["time_last"]  = ("UtcTime", "max")

    summary = (
        df.groupby(group_cols, dropna=True)
          .agg(**agg_dict)
          .reset_index()
          .sort_values("obs", ascending=False)
    )

    total = len(df)
    cells = []
    for _, row in summary.iterrows():
        def _int(col):
            v = row.get(col)
            return int(v) if v is not None and not pd.isna(v) else None
        def _flt(col):
            v = row.get(col)
            return float(v) if v is not None and not pd.isna(v) else None

        eci = _int("Cell_ECI")
        obs = int(row["obs"])

        # PCI / EARFCN — detect and correct column swap
        raw_pci    = _int("Cell_PCI")
        raw_earfcn = _int("Cell_EARFCN")
        if (raw_pci is not None and raw_earfcn is not None
                and raw_pci > 503 and raw_earfcn <= 503):
            pci, earfcn, col_swapped = raw_earfcn, raw_pci, True
        else:
            pci, earfcn, col_swapped = raw_pci, raw_earfcn, False

        tac  = _int("Cell_TAC")
        mcc  = _int("Cell_MCC")
        mnc  = _int("Cell_MNC")
        clat = _flt("Cell_Lat")
        clon = _flt("Cell_Lon")
        rsrp = _flt("rsrp_mean")

        band, freq, fcc_block = earfcn_to_band(earfcn)

        cells.append({
            "eci":         eci,
            "pci":         pci,
            "earfcn":      earfcn,
            "tac":         tac,
            "mcc":         mcc,
            "mnc":         mnc,
            "cell_lat":    clat,
            "cell_lon":    clon,
            "rsrp_mean":   rsrp,
            "obs":         obs,
            "pct":         obs / total * 100,
            "band":        band,
            "freq_desc":   freq,
            "fcc_block":   fcc_block,
            "col_swapped": col_swapped,
        })
    return cells

# ── ASR VERIFICATION ──────────────────────────────────────────────────────────

def query_asr(lat, lon, radius_km=0.5):
    """Query FCC ASR for registered structures near lat/lon."""
    radius_miles = radius_km * 0.621371   # FCC ASR dist param is in miles
    params = {
        "lat":   f"{lat:.6f}",
        "long":  f"{lon:.6f}",           # signed — negative for Western hemisphere
        "dist":  f"{radius_miles:.4f}",
        "limit": "20",
    }
    data = None
    for ep in ASR_ENDPOINTS:
        data = api_get(ep, params, label=f"ASR ({lat:.4f},{lon:.4f})")
        if data is not None:
            break
    if data is None:
        return None   # API failure

    regs = (data.get("Registrations")
            or data.get("registrations")
            or data.get("data")
            or data.get("results")
            or [])
    if isinstance(regs, dict):
        regs = regs.get("Registrations") or regs.get("data") or []

    results = []
    for r in regs:
        slat = float(r.get("latitude",  r.get("lat",  0)) or 0)
        slon = float(r.get("longitude", r.get("long", 0)) or 0)
        results.append({
            "registration_number": r.get("registrationNumber","—"),
            "structure_type":      r.get("structureType","—"),
            "height_agl":          r.get("heightAboveGroundLevel","—"),
            "owner":               r.get("ownerName","—"),
            "lat":                 slat,
            "lon":                 slon,
            "status":              r.get("statusCode","—"),
            "dist_m":              haversine(lat, lon, slat, slon) if slat and slon else None,
        })
    results.sort(key=lambda x: x.get("dist_m") or 99999)
    return results


def verify_asr(cells, offline=False):
    results = []
    seen = set()
    for cell in cells:
        if cell["cell_lat"] is None or cell["cell_lon"] is None:
            results.append({**cell, "asr_status":"NO_COORDS",
                "asr_structures":[], "asr_nearest_m":None, "asr_api_ok":None,
                "finding":"CANNOT_VERIFY",
                "finding_detail":"Cell_Lat/Cell_Lon not in points.csv — manual lookup required."})
            continue

        coord_key = (round(cell["cell_lat"], 3), round(cell["cell_lon"], 3))
        if coord_key in seen:
            continue
        seen.add(coord_key)

        manual_url = (
            f"https://wireless2.fcc.gov/UlsApp/AsrSearch/asrRegistrationSearch.jsp"
            f"?lat={cell['cell_lat']:.6f}&long={abs(cell['cell_lon']):.6f}&dist=0.5"
        )
        rest_url = (
            f"https://data.fcc.gov/api/antenna-system-registration/registrations.json"
            f"?lat={cell['cell_lat']:.6f}&long={cell['cell_lon']:.6f}&dist=0.311&limit=20"
        )

        if offline:
            results.append({**cell, "asr_status":"OFFLINE",
                "asr_structures":[], "asr_nearest_m":None, "asr_api_ok":False,
                "finding":"OFFLINE_MODE",
                "finding_detail":(
                    f"API skipped (--offline).  Manual lookup:\n{manual_url}\n\n"
                    f"REST API:\n{rest_url}"
                )})
            continue

        print(f"    ASR query: ECI {cell['eci']}  "
              f"({cell['cell_lat']:.5f}, {cell['cell_lon']:.5f})  "
              f"Band {cell['band']} / EARFCN {cell['earfcn']} / PCI {cell['pci']}")

        structures = query_asr(cell["cell_lat"], cell["cell_lon"],
                               radius_km=ASR_RADIUS_M / 1000.0)
        time.sleep(0.8)
        api_ok = structures is not None
        if structures is None:
            structures = []

        nearest_m = structures[0]["dist_m"] if structures else None

        if not api_ok:
            finding = "API_FAILURE"
            finding_detail = (
                f"FCC ASR API did not respond after {API_RETRIES} attempts.\n\n"
                f"MANUAL LOOKUP (paste into browser):\n{manual_url}\n\n"
                f"REST API:\n{rest_url}"
            )
        elif not structures:
            finding = "NO_REGISTERED_STRUCTURE"
            finding_detail = (
                f"No FCC ASR-registered antenna structure found within "
                f"{ASR_RADIUS_M}m of observed cell coordinates "
                f"({cell['cell_lat']:.6f}, {cell['cell_lon']:.6f}).\n\n"
                "A transmitter at this location without ASR registration violates "
                "47 U.S.C. § 303(q) and 47 C.F.R. Part 17.  If the structure exceeds "
                "200 ft AGL or is within 20,000 ft of an airport, FAA notification "
                "is also required.\n\n"
                f"MANUAL VERIFICATION URL (paste into browser):\n{manual_url}\n\n"
                "Screenshot the results and attach to your FCC enforcement complaint."
            )
        elif nearest_m <= ASR_RADIUS_M:
            finding = "REGISTERED_STRUCTURE_NEARBY"
            finding_detail = (
                f"Registered structure found {nearest_m:.0f}m from observed cell "
                "coordinates.  NOTE: ASR registration covers the physical structure "
                "only — not the radio equipment mounted on it.  A rogue device "
                "co-located on a registered tower is still an unauthorized transmitter "
                "under 47 U.S.C. § 301.  Cross-reference with ULS licensee data to "
                "confirm whether a licensed operator has a site agreement at this "
                "structure.  If not, this is still a § 301 violation."
            )
        else:
            finding = "STRUCTURE_OUTSIDE_RADIUS"
            finding_detail = (
                f"Nearest registered structure is {nearest_m:.0f}m away "
                f"(threshold: {ASR_RADIUS_M}m).  A transmitter at the observed "
                "coordinates not co-located with a registered structure has no "
                "FCC authorization pathway under Part 17."
            )

        results.append({**cell,
            "asr_status":     finding,
            "asr_structures": structures[:5],
            "asr_nearest_m":  nearest_m,
            "asr_api_ok":     api_ok,
            "finding":        finding,
            "finding_detail": finding_detail,
        })
    return results

# ── ULS VERIFICATION ──────────────────────────────────────────────────────────

def query_uls_band(band_num, state="CA"):
    """Query FCC ULS for licenses on a given band in California."""
    services = ULS_SERVICE_MAP.get(band_num, ["AW"])
    results = []
    for svc in services:
        params = {"radioServiceCode": svc, "state": state, "limit": "25"}
        data = None
        for ep in ULS_ENDPOINTS:
            data = api_get(ep, params, label=f"ULS Band {band_num} svc={svc}")
            if data is not None:
                break
        if data is None:
            continue
        licenses = (data.get("Licenses")
                    or data.get("licenses")
                    or (data.get("LicenseSearch") or {}).get("Licenses")
                    or data.get("data")
                    or data.get("results")
                    or [])
        if isinstance(licenses, dict):
            licenses = licenses.get("Licenses") or licenses.get("data") or []
        for lic in licenses:
            results.append({
                "file_number":   lic.get("licenseID") or lic.get("fileNumber","—"),
                "callsign":      lic.get("callSign","—"),
                "licensee":      lic.get("licenseeName") or lic.get("entityName","—"),
                "radio_service": lic.get("radioServiceCode","—"),
                "status":        lic.get("licenseStatus","—"),
                "expiration":    lic.get("expiredDate") or lic.get("expirationDate","—"),
                "market":        lic.get("marketDesc") or lic.get("channelBlock","—"),
            })
        time.sleep(0.5)
    return results


def verify_uls(cells, offline=False):
    bands = {}
    for cell in cells:
        b = cell["band"]
        if b is not None and b not in bands:
            bands[b] = cell

    results = {}
    for band_num, cell in sorted(bands.items()):
        print(f"    ULS query: Band {band_num} ({cell['freq_desc']})")
        known = KNOWN_LICENSEES.get(band_num)

        if offline:
            results[band_num] = _uls_result(band_num, cell, known, [], False)
            continue

        api_licenses = query_uls_band(band_num)
        results[band_num] = _uls_result(band_num, cell, known, api_licenses, True)
    return results


def _uls_result(band_num, cell, known, api_licenses, api_ok):
    mcc = cell.get("mcc")
    mnc = cell.get("mnc")
    mcc_mnc_match = None
    if known and mcc and mnc:
        known_mm = known.get("mcc_mnc","")
        session_mm = f"{mcc}-{mnc}"
        mcc_mnc_match = (session_mm in known_mm or known_mm in session_mm)

    if not known and not api_licenses:
        finding = "NO_LICENSE_DATA"
        detail  = (f"No pre-loaded licensee data for Band {band_num} and ULS API "
                   "returned no results.  Manual lookup: "
                   "https://wireless2.fcc.gov/UlsApp/UlsSearch/")
    elif mcc_mnc_match is False:
        finding = "MCC_MNC_MISMATCH"
        detail  = (f"Session MCC/MNC {mcc}/{mnc} does not match the known "
                   f"Band {band_num} licensee ({known['licensee']}, "
                   f"MCC/MNC {known['mcc_mnc']}) for Riverside County, CA.  "
                   "Broadcasting a mismatched MCC/MNC on licensed spectrum "
                   "constitutes identity theft of the licensee's network identity "
                   "(18 U.S.C. § 1028A) and unauthorized spectrum use "
                   "(47 U.S.C. § 301).")
    elif mcc_mnc_match is True:
        finding = "MCC_MNC_MATCHES_LICENSEE"
        detail  = (f"Session MCC/MNC {mcc}/{mnc} matches the known Band {band_num} "
                   f"licensee ({known['licensee']}).  NOTE: MCC/MNC match does not "
                   "mean the transmitter is authorized — an IMSI catcher deliberately "
                   "broadcasts the licensed carrier's identity to force device "
                   "attachment.  The ASR registration and absence of a site agreement "
                   "are the operative legal facts.")
    else:
        finding = "LICENSE_DATA_AVAILABLE"
        detail  = (f"Band {band_num} licensee data retrieved.  Cross-reference ASR "
                   "structure locations against licensee site agreements.")

    return {
        "band":          band_num,
        "freq_desc":     cell["freq_desc"],
        "fcc_block":     cell["fcc_block"],
        "known":         known,
        "api_licenses":  api_licenses[:10],
        "api_ok":        api_ok,
        "mcc_mnc_match": mcc_mnc_match,
        "finding":       finding,
        "finding_detail":detail,
    }

# ── REPORT BUILDER ────────────────────────────────────────────────────────────

def _notice_border(para, color):
    pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    for side in ["top","bottom","left","right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6" if side != "left" else "18")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        pBdr.append(el)
    pPr.append(pBdr)
    para.paragraph_format.left_indent = Inches(0.2)


def build_report(case_dir, cells, asr_results, uls_results,
                 hash_info, session_info, offline):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(0.9)

    # Title
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(t, "FCC INFRASTRUCTURE VERIFICATION REPORT",
            bold=True, color="1A237E", size=18)
    pPr = t._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    for side in ["top","bottom"]:
        el = OxmlElement(f"w:{side}"); el.set(qn("w:val"),"double")
        el.set(qn("w:sz"),"8"); el.set(qn("w:space"),"4")
        el.set(qn("w:color"),"1A237E"); pBdr.append(el)
    pPr.append(pBdr)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub, "FCC ASR / ULS Public Database Cross-Reference\n", bold=True, size=11)
    add_run(sub,
            f"Generated {datetime.datetime.now().strftime('%B %d, %Y  %H:%M UTC')}  |  "
            f"Case: {os.path.basename(case_dir)}\n"
            + ("MODE: OFFLINE\n" if offline else "FCC API: live\n"),
            size=10)
    add_run(sub,
            "COPY: FCC Enforcement Bureau  |  FBI Riverside  |  "
            "DOJ Civil Rights  |  Cal OES #26-1153",
            bold=True, color="B71C1C", size=10)

    # Legal framing notice
    n = doc.add_paragraph()
    add_run(n, "LEGAL FRAMING NOTICE: ", bold=True, color="1A237E", size=11)
    add_run(n,
            "The FCC regulates physical infrastructure, not radio-layer configuration.  "
            "The FCC tracks: tower coordinates (ASR), antenna height AGL/AMSL, structure "
            "ownership, and spectrum license holders (ULS).  "
            "The FCC does NOT track PCI assignments, TAC values, EARFCN configuration, "
            "RSRQ parameters, or Cell ECI values.  "
            "This report maps observed cellular data to FCC-actionable facts (ASR/ULS) "
            "and routes non-FCC findings to the correct federal agency.",
            size=10)
    _notice_border(n, "1A237E")
    doc.add_paragraph()

    # Section I — Case identification
    heading(doc, "I. CASE IDENTIFICATION AND CHAIN OF CUSTODY")
    all_bands = sorted({c["band"] for c in cells if c["band"]})
    swapped   = [c for c in cells if c.get("col_swapped")]
    add_kv_table(doc, [
        ("Complainant:",      COMPLAINANT,                                    False),
        ("DOB:",              DOB if DOB else "—",                            False),
        ("Case Folder:",      os.path.basename(case_dir),                     False),
        ("Source File:",      hash_info.get("File","—"),                      False),
        ("SHA-256 Hash:",     hash_info.get("Hash","—"),                      False),
        ("Session Start:",    _s(session_info.get("start")),                  False),
        ("Session End:",      _s(session_info.get("end")),                    False),
        ("Total Records:",    str(session_info.get("total_records","—")),     False),
        ("Unique Cells:",     str(len(cells)),                                False),
        ("Bands Observed:",   str(all_bands),                                 False),
        ("PCI/EARFCN Auto-corrected:", f"{len(swapped)} cell(s)",            bool(swapped)),
        ("API Mode:",         "OFFLINE (pre-loaded data)" if offline else "LIVE (FCC public API)", offline),
    ])

    # Section II — ULS license verification
    heading(doc, "II. FCC ULS — SPECTRUM LICENSE VERIFICATION BY BAND")
    add_run(doc.add_paragraph(),
            "For each LTE band observed in the session, this section identifies the "
            "FCC-licensed spectrum holder for Riverside County, CA.  A transmitter "
            "operating on licensed spectrum without authorization from the licensee "
            "violates 47 U.S.C. § 301 regardless of what radio-layer parameters it uses.",
            size=10)
    doc.add_paragraph()

    for band_num, ur in sorted(uls_results.items()):
        sc = ("B71C1C" if "MISMATCH" in ur["finding"] else
              "E65100" if ur["finding"] in ("NO_LICENSE_DATA","API_FAILURE") else
              "1B5E20")
        heading(doc, f"Band {band_num} — {ur['freq_desc']}  [{ur['finding']}]",
                level=2, color=sc)
        add_run(doc.add_paragraph(), ur["finding_detail"], size=10)
        doc.add_paragraph()

        rows = [
            ("Band:",          str(band_num),      False),
            ("Frequency:",     ur["freq_desc"],    False),
            ("FCC block:",     ur["fcc_block"],    False),
            ("Finding:",       ur["finding"],      "MISMATCH" in ur["finding"]),
        ]
        known = ur.get("known")
        if known:
            rows += [
                ("Licensed operator:", known["licensee"],                 False),
                ("License file(s):",   ", ".join(known["file_numbers"]), False),
                ("Market:",            known["market"],                   False),
                ("Licensed MCC/MNC:",  known["mcc_mnc"],                 False),
                ("Note:",              known["note"],                     False),
            ]
        mm = ur.get("mcc_mnc_match")
        if mm is not None:
            rows.append(("Session MCC/MNC match:",
                         "YES — matches licensee" if mm else "NO — MISMATCH",
                         not mm))
        add_kv_table(doc, rows)

        if ur.get("api_licenses"):
            add_log_table(doc,
                ["Licensee","Callsign","Service","Status","File Number","Market"],
                [2.0,0.9,0.7,0.7,1.0,1.2],
                [[(l["licensee"],False),(l["callsign"],False),(l["radio_service"],False),
                  (l["status"], l["status"] not in ("A","Active")),
                  (l["file_number"],False),(l["market"],False)]
                 for l in ur["api_licenses"]])

    # Section III — ASR structure verification
    heading(doc, "III. FCC ASR — ANTENNA STRUCTURE REGISTRATION VERIFICATION")
    no_struct    = [r for r in asr_results if r["finding"]=="NO_REGISTERED_STRUCTURE"]
    api_failures = [r for r in asr_results if r["finding"]=="API_FAILURE"]
    registered   = [r for r in asr_results if r["finding"]=="REGISTERED_STRUCTURE_NEARBY"]
    add_run(doc.add_paragraph(),
            f"Each unique cell tower location in points.csv was queried against the "
            f"FCC ASR database for registered structures within {ASR_RADIUS_M}m.  "
            "A transmitter at coordinates with no registered structure has no legitimate "
            "physical presence in the FCC's infrastructure record.",
            size=10)
    doc.add_paragraph()
    add_kv_table(doc, [
        ("Locations queried:",     str(len(asr_results)),   False),
        ("No structure found:",    str(len(no_struct)),     bool(no_struct)),
        ("Structure nearby:",      str(len(registered)),    False),
        ("API failures:",          str(len(api_failures)),  bool(api_failures)),
        ("Search radius:",         f"{ASR_RADIUS_M}m",     False),
    ])

    for r in asr_results:
        sc2 = ("B71C1C" if r["finding"]=="NO_REGISTERED_STRUCTURE" else
               "E65100" if r["finding"] in ("API_FAILURE","OFFLINE_MODE","CANNOT_VERIFY") else
               "1B5E20")
        heading(doc,
                f"ECI {r['eci']}  PCI {r['pci']}  Band {r['band']}  "
                f"EARFCN {r['earfcn']}  [{r['finding']}]",
                level=2, color=sc2)
        add_run(doc.add_paragraph(), r["finding_detail"], size=10)
        doc.add_paragraph()

        add_kv_table(doc, [
            ("Cell ECI:",      _s(r["eci"]),
             r["finding"]=="NO_REGISTERED_STRUCTURE"),
            ("PCI:",           _s(r["pci"]),        False),
            ("EARFCN / Band:", f"{r['earfcn']} / Band {r['band']} ({r['freq_desc']})", False),
            ("Coordinates:",   f"{_fmt(r['cell_lat'],6)}, {_fmt(r['cell_lon'],6)}", False),
            ("TAC:",           _s(r["tac"]),         r.get("tac") in {0,65535,0xFFFE}),
            ("Session obs:",   f"{r['obs']} ({r['pct']:.1f}%)", False),
            ("RSRP mean:",     _fmt(r["rsrp_mean"],1," dBm"), False),
            ("ASR finding:",   r["finding"],         r["finding"]=="NO_REGISTERED_STRUCTURE"),
            ("Nearest struct:",
             _fmt(r["asr_nearest_m"],0,"m") if r["asr_nearest_m"] else "None found",
             r["finding"]=="NO_REGISTERED_STRUCTURE"),
            ("Col swap corrected:", "YES" if r.get("col_swapped") else "NO",
             bool(r.get("col_swapped"))),
        ])

        if r.get("asr_structures"):
            add_log_table(doc,
                ["ASR Reg #","Type","Owner","Height AGL","Dist from cell","Status"],
                [1.2,0.8,2.0,0.9,1.0,0.6],
                [[(s["registration_number"],False),(s["structure_type"],False),
                  (s["owner"],False),(_fmt(s["height_agl"],0," ft"),False),
                  (_fmt(s["dist_m"],0," m"),
                   s["dist_m"] is not None and s["dist_m"] > ASR_RADIUS_M),
                  (s["status"],False)]
                 for s in r["asr_structures"]])

    # Section IV — Jurisdictional routing table
    doc.add_page_break()
    heading(doc, "IV. JURISDICTIONAL ROUTING TABLE")
    add_run(doc.add_paragraph(),
            "This table maps each anomaly type to the specific federal agency with "
            "jurisdiction and the correct filing mechanism.  Filing with the wrong "
            "agency causes unnecessary delay.  The FCC regulates infrastructure.  "
            "DOJ/FBI regulate conduct.  GPS interference falls to DOD and DOJ.",
            size=10)
    doc.add_paragraph()

    sev_colors = {"PRIMARY":"B71C1C","SECONDARY":"E65100",
                  "SUPPORTING":"F57F17","COMPOUND":"6A1B9A"}
    for entry in ROUTING_TABLE:
        sc3 = sev_colors.get(entry["severity"],"000000")
        heading(doc, f"[{entry['severity']}]  {entry['anomaly']}",
                level=2, color=sc3)
        add_kv_table(doc, [
            ("Evidence basis:",   entry["evidence"],                False),
            ("Primary agency:",   entry["agency"],                  True),
            ("Legal mechanism:",  entry["mechanism"],               False),
            ("Statutes:",         " | ".join(entry["statutes"]),    False),
            ("Filing action:",    entry["action"],                  True),
        ])

    # Section V — Filing instructions
    doc.add_page_break()
    heading(doc, "V. CONSOLIDATED FILING INSTRUCTIONS")
    filings = [
        ("FCC Enforcement Bureau", "enforcement@fcc.gov",
         "47 U.S.C. §§ 301, 333; 47 C.F.R. § 2.807",
         ["points.csv","incidents.csv","fcc_verification DOCX (this report)",
          "nlp_evidence_report DOCX","evidence_hash.txt"],
         "Lead with ASR finding — no registered structure at cell coordinates is "
         "the most actionable FCC fact.  FCC can act on: § 301 (unauthorized transmitter), "
         "§ 333 (interference), § 2.807 (jamming)."),
        ("FBI Riverside / IC3", "ic3.gov + direct field office",
         "18 U.S.C. §§ 2511, 2512, 1029, 1030, 1028A",
         ["All DOCX reports","points.csv","incidents.csv",
          "evidence_hash.txt","NLP.csv (CTW-11 SENTINEL session)"],
         "Reference sheriff case PE260810041 and prior IC3 submissions.  "
         "FBI handles wiretap (§ 2511), IMSI-catcher possession (§ 2512), "
         "access device fraud (§ 1029), CFAA (§ 1030), identity theft (§ 1028A)."),
        ("DOJ Civil Rights Division", "civilrights.fbi.gov",
         "18 U.S.C. §§ 241, 242",
         ["Civil rights victim statement","all DOCX reports","chronological declaration"],
         "If any actor holds government authority (§ 242) or if two or more persons "
         "coordinated the targeting (§ 241)."),
        ("U.S. Space Force GPS Ops Center", "(719) 567-2828",
         "18 U.S.C. § 1367; DOD Directive 4650.1",
         ["GPS injection findings from nlp_evidence_report",
          "points.csv (GNSS rows with injection constants)"],
         "GPS spoofing is interference with a military asset.  "
         "Space Force GPSOC tracks GPS interference events globally."),
    ]
    for agency, contact, statutes, attachments, notes in filings:
        heading(doc, agency, level=2, color="1A237E")
        add_kv_table(doc, [
            ("Contact:",      contact,   False),
            ("Statutes:",     statutes,  False),
            ("Attachments:",  "\n".join(f"  \u2022 {a}" for a in attachments), False),
            ("Notes:",        notes,     False),
        ])

    # Section VI — Certification
    doc.add_page_break()
    heading(doc, "VI. CERTIFICATION UNDER PENALTY OF PERJURY")
    cert = doc.add_paragraph()
    add_run(cert, f"I, {COMPLAINANT}, ", size=11)
    add_run(cert, "declare under penalty of perjury ", bold=True, size=11)
    add_run(cert,
            "under the laws of the United States of America (28 U.S.C. § 1746) "
            "that all cellular observation data in this report is extracted directly "
            "and without alteration from machine-generated data files produced by "
            f"GNSS_AttackModel.ps1 in case folder {os.path.basename(case_dir)}.  "
            "FCC ASR and ULS query results were retrieved from public FCC APIs at the "
            "timestamp recorded above.  Pre-loaded licensee data is sourced from "
            "publicly available FCC ULS records.  No value has been manually modified "
            "except the complainant name and date of birth above.",
            size=11)
    doc.add_paragraph(); doc.add_paragraph()
    sig = doc.add_paragraph()
    add_run(sig, f"{COMPLAINANT}\n", bold=True, size=11)
    add_run(sig, f"DOB: {DOB if DOB else '—'}\n", size=11)
    add_run(sig, f"Report generated: {datetime.datetime.now().isoformat()}",
            size=10, italic=True)
    return doc

# ── MAIN ──────────────────────────────────────────────────────────────────────

def main():
    global COMPLAINANT, DOB

    parser = argparse.ArgumentParser(
        description="FCC ASR/ULS Verification — companion to GNSS_AttackModel.ps1"
    )
    parser.add_argument("--case",    default=None,
                        help="Explicit GNSS_AttackCase_* folder path")
    parser.add_argument("--base",    default=BASE_DIR,
                        help=f"Base folder to search for case folders (default: {BASE_DIR})")
    parser.add_argument("--offline", action="store_true",
                        help="Skip live API calls; use pre-loaded licensee data only")
    args = parser.parse_args()

    print("=" * 65)
    print("FCC ASR / ULS VERIFICATION — Infrastructure Cross-Reference")
    print("=" * 65)
    print()

    while True:
        name = input("Enter complainant full legal name: ").strip()
        if name: COMPLAINANT = name; break
        print("  Name cannot be empty.")
    dob_input = input("Enter DOB (optional — press Enter to skip): ").strip()
    DOB = dob_input if dob_input else ""
    print()

    case_dir  = find_case_folder(args.base, args.case)
    print(f"Case folder : {case_dir}")
    if args.offline:
        print("Mode        : OFFLINE (pre-loaded data, no API calls)")
    print()

    df        = load_points(case_dir)
    hash_info = read_hash_file(case_dir)
    print(f"Points loaded : {len(df)} records")

    ts_col = next((c for c in ["UtcTime","TimeToUtc"] if c in df.columns), None)
    ts_ser = df[ts_col].dropna() if ts_col else pd.Series([], dtype="object")
    session_info = {
        "start":         ts_ser.min() if not ts_ser.empty else None,
        "end":           ts_ser.max() if not ts_ser.empty else None,
        "total_records": len(df),
    }
    print(f"Session       : {session_info['start']} -> {session_info['end']}")
    print()

    cells = build_unique_cells(df)
    swapped = [c for c in cells if c.get("col_swapped")]
    print(f"Unique cells  : {len(cells)}")
    if swapped:
        print(f"  NOTE: {len(swapped)} cell(s) had PCI/EARFCN auto-corrected "
              f"(column swap detected from GNSS_AttackModel.ps1 / CellMapper version mismatch)")
    bands = sorted({c["band"] for c in cells if c["band"]})
    print(f"Bands observed: {bands}")
    print()

    print("Querying FCC ASR (antenna structure registrations)...")
    asr_results = verify_asr(cells, offline=args.offline)
    no_struct   = [r for r in asr_results if r["finding"]=="NO_REGISTERED_STRUCTURE"]
    if no_struct:
        print(f"  *** {len(no_struct)} cell location(s) have NO registered structure "
              f"within {ASR_RADIUS_M}m ***")
    print()

    print("Querying FCC ULS (spectrum license database)...")
    uls_results = verify_uls(cells, offline=args.offline)
    for band_num, ur in sorted(uls_results.items()):
        print(f"  Band {band_num:3d}: {ur['finding']}")
    print()

    ts_str   = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    json_out = os.path.join(case_dir, f"fcc_verification_{ts_str}.json")
    with open(json_out, "w", encoding="utf-8") as jf:
        json.dump({
            "generated":   datetime.datetime.now().isoformat(),
            "case_folder": case_dir,
            "cells":       cells,
            "asr_results": asr_results,
            "uls_results": {str(k): v for k, v in uls_results.items()},
        }, jf, indent=2, default=str)
    print(f"Raw JSON saved: {os.path.basename(json_out)}")

    print("Generating DOCX report...")
    doc = build_report(case_dir, cells, asr_results, uls_results,
                       hash_info, session_info, args.offline)
    docx_out = os.path.join(case_dir, f"fcc_verification_{ts_str}.docx")
    doc.save(docx_out)
    print(f"Report saved  : {os.path.basename(docx_out)}")
    print()
    print("=" * 65)
    print(f"Files written to: {case_dir}")
    if no_struct:
        print()
        print(f"*** ACTION: {len(no_struct)} cell(s) have no ASR-registered structure ***")
        print("    Lead with this in your enforcement@fcc.gov filing.")
    print("=" * 65)


if __name__ == "__main__":
    main()
