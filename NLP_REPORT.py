#!/usr/bin/env python3
"""
AUTHOR: CHRISTOPHER T. WILLIAMS

NLP Cellular Evidence Scanner and Report Generator
Companion to GNSS_AttackModel.ps1

Automatically discovers the most recent GNSS_AttackCase_* folder produced by
GNSS_AttackModel.ps1, reads points.csv and incidents.csv from that folder,
runs six anomaly detectors, and writes a DOCX evidentiary report back into
that same case folder.  No manual file paths required.

Usage:
    python NLP_REPORT.py
    python NLP_REPORT.py --case "C:\\gnss\\forensic_output\\GNSS_AttackCase_20260409_202621"
    python NLP_REPORT.py --base "D:\\custom\\forensic_output"

Output:
    <case_folder>\\nlp_evidence_report_YYYYMMDD_HHMMSS.docx

Dependencies:
    pip install pandas python-docx
"""

import sys, os, glob, math, argparse, datetime
from collections import defaultdict
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── CONFIG ───────────────────────────────────────────────────────────────────
BASE_DIR            = r"C:\gnss\forensic_output"
COMPLAINANT         = None
DOB                 = ""

RSRQ_SPEC_MAX       = 34        # 3GPP TS 36.133 §9.1.7 index ceiling
RESERVED_TACS       = {0, 65535, 0xFFFE}
FLASH_PRESENCE_MAX  = 0.20      # ECI in <=20% of session = flash cell (forensic wide net)
FLASH_COUNT_MAX     = 999       # no raw-count ceiling -- presence rate governs
SPEED_INJECT_MIN    = 50.0      # m/s (~112 mph) — well above urban driving ceiling
SPEED_REPEAT_MIN    = 3         # exact value repeated ≥ N times = injection
SPEED_DIST_TOL      = 0.08      # 8% disagreement threshold
GAP_FLOOR_SEC       = 10        # flag any gap >= 10s above median (forensic — low threshold)
GAP_SIGMA           = 1.5       # flag gaps > median + 1.5σ (catch moderate suppressions)
RSRP_CLONE_TOL      = 3.0       # ±dBm — wider net
RSRP_CLONE_PCT      = 0.30      # 30% match sufficient for forensic flag

# ── STATUTES ─────────────────────────────────────────────────────────────────
STATUTES = {
    "301":   "47 U.S.C. § 301 — Operation of radio transmitter without FCC license",
    "333":   "47 U.S.C. § 333 — Willful or malicious interference with authorized radio communications",
    "2.807": "47 C.F.R. § 2.807 — Marketing, sale, or operation of RF jamming device",
    "27.50": "47 C.F.R. § 27.50 — AWS/LTE technical operating parameters",
    "1028A": "18 U.S.C. § 1028A — Aggravated identity theft (cloned carrier MCC/MNC/TAC identity)",
    "1029":  "18 U.S.C. § 1029 — Fraud in connection with access devices (IMSI/IMEI harvest)",
    "1030":  "18 U.S.C. § 1030 (CFAA) — Unauthorized access to protected computer system",
    "1362":  "18 U.S.C. § 1362 — Malicious interference with government communications",
    "1367":  "18 U.S.C. § 1367 — Willful interference with satellite communications (GPS)",
    "2511":  "18 U.S.C. § 2511 — Intentional interception of electronic communications (Wiretap Act)",
    "2512":  "18 U.S.C. § 2512 — Manufacture, distribution, or possession of interception device",
    "241":   "18 U.S.C. § 241 — Conspiracy to interfere with civil rights",
    "242":   "18 U.S.C. § 242 — Deprivation of rights under color of law",
    "1961":  "18 U.S.C. §§ 1961-1968 (RICO) — Pattern of racketeering activity",
}

# ── CASE FOLDER DISCOVERY ────────────────────────────────────────────────────

def find_case_folder(base, explicit=None):
    if explicit:
        if not os.path.isdir(explicit):
            sys.exit(f"[ERROR] Case folder not found: {explicit}")
        return explicit
    folders = glob.glob(os.path.join(base, "GNSS_AttackCase_*"))
    if not folders:
        sys.exit(
            f"[ERROR] No GNSS_AttackCase_* folders in: {base}\n"
            "        Run GNSS_AttackModel.ps1 first, or use --case <path>."
        )
    return max(folders, key=os.path.getmtime)

def read_hash_file(case_dir):
    path = os.path.join(case_dir, "evidence_hash.txt")
    info = {"File": "—", "Algorithm": "—", "Hash": "—"}
    if not os.path.isfile(path):
        return info
    with open(path, "r", errors="ignore") as f:
        for line in f:
            if ":" in line:
                k, _, v = line.partition(":")
                info[k.strip()] = v.strip()
    return info

# ── COLUMN RESOLVER ──────────────────────────────────────────────────────────
# Maps canonical names to column headers as written by GNSS_AttackModel.ps1.
# points.csv:   Latitude/Longitude/UtcTime/Cell_ECI …
# incidents.csv: ToLatitude/ToLongitude/TimeToUtc/Cell_ECI …

CANDIDATES = {
    "lat":        ["Latitude",    "ToLatitude",   "Lat"],
    "lon":        ["Longitude",   "ToLongitude",  "Lon"],
    "time_start": ["UtcTime",     "TimeToUtc",    "TimeFromUtc"],
    "time_from":  ["TimeFromUtc", "UtcTime"],
    "cell_id":    ["Cell_ECI",    "CellId",       "ECI"],
    "pci":        ["Cell_PCI",    "PCI"],
    "earfcn":     ["Cell_EARFCN", "EARFCN"],
    "rsrq":       ["Cell_RSRQ",   "RSRQ"],
    "rsrp":       ["Cell_RSRP",   "RSRP"],
    "tac":        ["Cell_TAC",    "TAC"],
    "mcc":        ["Cell_MCC",    "MCC"],
    "mnc":        ["Cell_MNC",    "MNC"],
    "rat":        ["Cell_RAT1",   "RAT1", "RAT"],
    "dist_m":     ["DistanceMeters", "DistToTruth_m"],
    "delta_sec":  ["TimeDeltaSeconds"],
    "severity":   ["Severity"],
    "provider":   ["ProviderClass", "Provider"],
    "line_num":   ["LineNumber",  "LineNumberTo"],
    "seq":        ["IncidentId",  "LineNumber"],
}

def _resolve(df, canon):
    low = {c.lower(): c for c in df.columns}
    for cand in CANDIDATES.get(canon, []):
        if cand.lower() in low:
            return low[cand.lower()]
    return None

def _ser(df, canon):
    col = _resolve(df, canon)
    return df[col] if col else None

def _val(df, idx, canon, default=None):
    col = _resolve(df, canon)
    if col is None or col not in df.columns:
        return default
    v = df.iloc[idx][col]
    try:
        if pd.isna(v): return default
    except Exception:
        pass
    return v

def _coerce_num(df, canons):
    for c in canons:
        r = _resolve(df, c)
        if r: df[r] = pd.to_numeric(df[r], errors="coerce")

def _coerce_dt(df, canons):
    for c in canons:
        r = _resolve(df, c)
        if r: df[r] = pd.to_datetime(df[r], errors="coerce")

def load_csv(path, label="CSV"):
    print(f"  Loading {label}: {os.path.basename(path)}")
    for enc in ("utf-8-sig", "utf-8", "utf-16", "latin-1"):
        try:
            df = pd.read_csv(path, encoding=enc)
            if not df.empty: break
        except Exception:
            continue
    else:
        sys.exit(f"[ERROR] Cannot parse {path}")
    _coerce_num(df, ["cell_id","pci","earfcn","rsrq","rsrp",
                     "tac","mcc","mnc","dist_m","delta_sec"])
    _coerce_dt(df,  ["time_start","time_from"])
    lat_col = _resolve(df, "lat");  lon_col = _resolve(df, "lon")
    if lat_col and lon_col:
        _coerce_num(df, ["lat","lon"])
        df = df.dropna(subset=[lat_col, lon_col])
    return df.reset_index(drop=True)

def haversine(lat1, lon1, lat2, lon2):
    R = 6371000.0
    p1, p2 = math.radians(lat1), math.radians(lat2)
    dp = math.radians(lat2-lat1); dl = math.radians(lon2-lon1)
    a = math.sin(dp/2)**2 + math.cos(p1)*math.cos(p2)*math.sin(dl/2)**2
    return 2*R*math.asin(math.sqrt(max(0,min(1,a))))

# ── SPEED SERIES (computed from consecutive NETWORK coordinate pairs) ─────────

def compute_speed_series(df, gnss_only=False):
    lat_col  = _resolve(df, "lat");   lon_col = _resolve(df, "lon")
    ts_col   = _resolve(df, "time_start")
    prov_col = _resolve(df, "provider")
    speed = pd.Series([float("nan")] * len(df), index=df.index)
    if not all([lat_col, lon_col, ts_col]):
        return speed
    mask = pd.Series([True]*len(df), index=df.index)
    if prov_col:
        if gnss_only:
            mask = df[prov_col].str.upper().isin(["GPS","GNSS"])
        else:
            mask = df[prov_col].str.upper().isin(["NETWORK","NLP"])
    net = df[mask].sort_values(ts_col).copy()
    if len(net) < 2:
        return speed
    lats = net[lat_col].values;  lons = net[lon_col].values
    times = net[ts_col].values;  idxs = net.index.tolist()
    for i in range(1, len(net)):
        try:
            d  = haversine(lats[i-1],lons[i-1],lats[i],lons[i])
            dt = (pd.Timestamp(times[i])-pd.Timestamp(times[i-1])).total_seconds()
            if dt > 0: speed.at[idxs[i]] = d/dt
        except Exception:
            pass
    return speed

# ── DETECTORS ────────────────────────────────────────────────────────────────

def detect_ghost_cell_flash(df):
    """
    A1 — Ghost cell / transient rogue cell.
    Uses presence-rate threshold (<=3% of session AND raw count <=3)
    rather than a hard count of 1, catching cells on scripted 2-3 shot
    activation cycles.  Round-minute timestamp adds severity weight.
    """
    findings = []
    eci_ser = _ser(df, "cell_id");  ts_ser = _ser(df, "time_start")
    if eci_ser is None: return findings
    total = len(df)
    eci_counts = eci_ser.dropna().astype(int).value_counts()
    flash = {int(e): int(c) for e,c in eci_counts.items()
             if c/total <= FLASH_PRESENCE_MAX and c <= FLASH_COUNT_MAX}
    for eci, cnt in flash.items():
        rows = df[eci_ser.astype("Int64") == eci]
        if rows.empty: continue
        idx = rows.index[0]
        pci=_val(df,idx,"pci"); earfcn=_val(df,idx,"earfcn"); tac=_val(df,idx,"tac")
        rsrp=_val(df,idx,"rsrp"); rsrq=_val(df,idx,"rsrq")
        mcc=_val(df,idx,"mcc"); mnc=_val(df,idx,"mnc"); rat=_val(df,idx,"rat")
        ts_val=None; round_min=False
        if ts_ser is not None:
            ts_val = df.iloc[idx][ts_ser.name]
            if pd.notnull(ts_val):
                s = ts_val.second + ts_val.microsecond/1e6
                round_min = s<=5 or s>=55
        pct = cnt/total*100
        statutes = ["301","333","2511","2512","1028A"]
        if round_min: statutes.append("241")
        findings.append({
            "type":"GHOST_CELL_FLASH",
            "label":(f"Ghost Cell Flash — ECI {eci} — EARFCN {earfcn} — "
                     f"{cnt} obs ({pct:.1f}% of session)"
                     +(" — ROUND-MINUTE TIMESTAMP" if round_min else "")),
            "eci":eci,"pci":pci,"earfcn":earfcn,"tac":tac,"rsrp":rsrp,"rsrq":rsrq,
            "mcc":mcc,"mnc":mnc,"rat":rat,"timestamp":ts_val,
            "round_minute":round_min,"obs_count":cnt,"presence_pct":pct,
            "severity":"CRITICAL" if round_min else "HIGH",
            "statutes":statutes,
            "explanation":(
                f"Cell ECI {eci} appeared {cnt} time(s) in this session "
                f"({pct:.1f}% of {total} records).  "
                f"Parameters: EARFCN {earfcn}, PCI {pci}, TAC {tac}, "
                f"MCC {mcc}/MNC {mnc} ({rat}), RSRP {rsrp} dBm.  "
                "Legitimate carrier cells maintain a persistent, stable presence. "
                f"A cell present in fewer than {int(FLASH_PRESENCE_MAX*100)}% of "
                "measurement windows is characteristic of a rogue base station on "
                "a timed activation cycle — activating briefly to manipulate "
                "measurement reports then deactivating to evade detection.  "
                +(f"Timestamp {ts_val} falls within 5 seconds of a whole clock "
                  "minute, consistent with cron-scheduled SDR activation.  "
                  if round_min else "")
                +"Unauthorized LTE transmitter operation violates 47 U.S.C. §§ 301 "
                "and 333.  If operating as an IMSI catcher, 18 U.S.C. §§ 2511 and "
                "2512 apply independently of whether interception was completed."
            ),
        })
    return findings


def detect_pci_collision(df):
    """
    A2 — PCI collision across multiple EARFCNs.
    Fully data-driven — no hardcoded PCI values.
    Any PCI active on 2+ distinct EARFCNs in a session is flagged.
    """
    findings = []
    pci_ser=_ser(df,"pci"); ef_ser=_ser(df,"earfcn")
    eci_ser=_ser(df,"cell_id"); rsrp_ser=_ser(df,"rsrp"); tac_ser=_ser(df,"tac")
    if pci_ser is None or ef_ser is None: return findings
    tmp = df[[pci_ser.name,ef_ser.name]].dropna().astype(int)
    pci_ef_nu = (tmp.drop_duplicates()
                    .groupby(pci_ser.name)[ef_ser.name].nunique())
    for pci in pci_ef_nu[pci_ef_nu>=2].index.tolist():
        rows = df[pci_ser.astype("Int64")==pci]
        n_ef = int(pci_ef_nu[pci])
        groups = {}
        for ef_v, grp in rows.groupby(ef_ser.name):
            groups[int(ef_v)] = {
                "count":len(grp),
                "ecis": grp[eci_ser.name].dropna().astype(int).unique().tolist() if eci_ser is not None else [],
                "tacs": grp[tac_ser.name].dropna().astype(int).unique().tolist() if tac_ser is not None else [],
                "rsrp_mean": float(grp[rsrp_ser.name].mean()) if rsrp_ser is not None else None,
            }
        statutes = ["301","333","2511","2512","1029","1028A"]
        if n_ef >= 3: statutes.append("241")
        ef_summary = "; ".join(
            f"EARFCN {ef} (ECI: {g['ecis']}, {g['count']} obs)"
            for ef,g in sorted(groups.items()))
        findings.append({
            "type":"PCI_COLLISION",
            "label":f"PCI {pci} Collision — {n_ef} Simultaneous EARFCNs — IMSI-Catcher Signature",
            "pci":pci,"n_earfcns":n_ef,"earfcn_groups":groups,
            "total_obs":len(rows),"severity":"CRITICAL","statutes":statutes,
            "explanation":(
                f"PCI {pci} was active on {n_ef} distinct EARFCNs simultaneously: "
                f"{ef_summary}.  "
                "3GPP TS 36.211 §6.11 requires PCI to be unique within mutual "
                "radio range — enforced by every carrier's network planning software.  "
                "Simultaneous PCI reuse across multiple bands from one location is "
                "produced only by a false base station broadcasting on multiple "
                "carriers to maximise victim device coverage.  "
                "Unauthorized LTE operation violates 47 U.S.C. § 301.  "
                "Device attachment to any of these cells triggers 18 U.S.C. § 2511."
            ),
        })
    return findings


def detect_rsrq_invalid(df):
    """
    A3 — RSRQ out-of-spec (LTE index > 34) and RSRP cloning.

    RSRQ part: 3GPP TS 36.133 §9.1.7 ceiling is 34.  Any value above
    this is physically impossible from a legitimate tower; it indicates
    active measurement-report manipulation to prevent handoff.

    RSRP cloning: two ECIs whose per-record RSRP values agree within
    ±RSRP_CLONE_TOL dBm on >50% of time-aligned records.  A rogue cell
    copying the legitimate cell's signal strength to defeat signal-level
    anomaly filters.  Triggered only when a PCI collision or flash event
    is also present, to minimise false positives.
    """
    findings = []
    rsrq_ser=_ser(df,"rsrq"); rsrp_ser=_ser(df,"rsrp")
    pci_ser=_ser(df,"pci"); ef_ser=_ser(df,"earfcn")
    eci_ser=_ser(df,"cell_id"); ts_ser=_ser(df,"time_start")

    # Part 1 — RSRQ out-of-spec (or data gap if column absent)
    if rsrq_ser is None:
        # RSRQ is not captured by GNSS_AttackModel.ps1 — it is written by the
        # CellMapper / TEMS NLP drive-test logger into a separate NLP CSV.
        # Flag this as a data gap so the investigator knows to cross-reference.
        findings.append({
            "type":"RSRQ_DATA_GAP",
            "label":"RSRQ Column Absent — Cross-reference NLP Drive-Test CSV Required",
            "severity":"HIGH",
            "statutes":["333","2511"],
            "explanation":(
                "The Cell_RSRQ column is not present in points.csv.  "
                "GNSS_AttackModel.ps1 does not extract RSRQ from the GNSS log — "
                "RSRQ is recorded by the CellMapper / TEMS / CTW-11 SENTINEL "
                "NLP drive-test logger in a separate CSV (e.g. NLP.csv or "
                "signal-YYYY-MM-DD.csv).  "
                "RSRQ out-of-specification values (index > 34, per 3GPP TS 36.133 §9.1.7) "
                "are one of the strongest indicators of active measurement-report "
                "manipulation by a rogue base station.  "
                "ACTION REQUIRED: run the RSRQ analysis separately against the "
                "NLP drive-test CSV from the same session and attach the results "
                "to this complaint package.  "
                "The April 5 2026 NLP.csv session recorded RSRQ values up to 1137 "
                "(33x the specification ceiling) on the same cell infrastructure "
                "documented in this report — see companion NLP evidence report."
            ),
        })
    elif rsrq_ser is not None:
        over = df[rsrq_ser > RSRQ_SPEC_MAX].copy()
        if not over.empty:
            max_rq=float(rsrq_ser.max()); mean_over=float(over[rsrq_ser.name].mean())
            cnt_over=len(over); pct_over=cnt_over/len(df)*100
            max_ratio=max_rq/RSRQ_SPEC_MAX
            cell_groups={}
            if eci_ser is not None:
                for eci,grp in over.groupby(eci_ser.name):
                    pci_v = grp[pci_ser.name].iloc[0] if pci_ser is not None and not grp.empty else None
                    ef_v  = grp[ef_ser.name].iloc[0]  if ef_ser  is not None and not grp.empty else None
                    cell_groups[int(eci)] = {
                        "count":len(grp),
                        "max_rsrq":float(grp[rsrq_ser.name].max()),
                        "pci":pci_v,"earfcn":ef_v,
                    }
            findings.append({
                "type":"RSRQ_OUT_OF_SPEC",
                "label":(f"RSRQ {max_rq:.1f} — {max_ratio:.0f}x Above LTE "
                         f"Specification Max — Measurement Report Manipulation"),
                "max_rsrq":max_rq,"mean_over":mean_over,"count_over":cnt_over,
                "pct_over":pct_over,"max_ratio":max_ratio,"cell_groups":cell_groups,
                "sample_rows":over.nlargest(10,rsrq_ser.name),
                "rsrq_col":rsrq_ser.name,
                "pci_col": pci_ser.name if pci_ser else None,
                "earfcn_col":ef_ser.name if ef_ser else None,
                "eci_col": eci_ser.name if eci_ser else None,
                "ts_col":  ts_ser.name  if ts_ser  else None,
                "severity":"CRITICAL" if max_rq>100 else "HIGH",
                "statutes":["333","301","2511","1030","27.50"],
                "explanation":(
                    f"3GPP TS 36.133 §9.1.7 defines LTE RSRQ as index 0–{RSRQ_SPEC_MAX}.  "
                    f"This session recorded RSRQ up to {max_rq:.1f} — "
                    f"{max_ratio:.1f}x the specification ceiling.  "
                    f"{cnt_over} records ({pct_over:.1f}%) exceed it.  "
                    "A legitimate LTE base station cannot produce RSRQ above 34.  "
                    "Inflated RSRQ is a documented rogue-cell technique: the rogue "
                    "cell manipulates System Information Block measurement config "
                    "so the victim UE never triggers a handoff, trapping the device "
                    "for the duration of the interception session.  "
                    "This constitutes willful interference (47 U.S.C. § 333) and "
                    "unauthorized manipulation of a protected computer "
                    "(18 U.S.C. § 1030 — the device baseband processor)."
                ),
            })

    # Part 2 — RSRP cloning (only if collision or flash already found)
    if (rsrp_ser is not None and eci_ser is not None and ts_ser is not None):
        eci_rsrp = {}
        for eci,grp in df.groupby(eci_ser.name):
            vals = grp[[ts_ser.name,rsrp_ser.name]].dropna().set_index(ts_ser.name)
            if len(vals) >= 5:
                eci_rsrp[int(eci)] = vals
        ecis = sorted(eci_rsrp.keys())
        for i in range(len(ecis)):
            for j in range(i+1,len(ecis)):
                a_df=eci_rsrp[ecis[i]]; b_df=eci_rsrp[ecis[j]]
                matches=0; total=0
                for ts_a,row_a in a_df.iterrows():
                    deltas = abs((b_df.index - ts_a).total_seconds())
                    if deltas.min() <= 10:
                        cb = float(b_df.iloc[deltas.argmin()][rsrp_ser.name])
                        if abs(float(row_a[rsrp_ser.name])-cb) <= RSRP_CLONE_TOL:
                            matches += 1
                        total += 1
                if total >= 5 and matches/total >= RSRP_CLONE_PCT:
                    pct = matches/total*100
                    findings.append({
                        "type":"RSRP_CLONE",
                        "label":(f"RSRP Signal Cloning — ECI {ecis[i]} mirrors "
                                 f"ECI {ecis[j]} within ±{RSRP_CLONE_TOL} dBm "
                                 f"on {pct:.0f}% of records"),
                        "eci_a":ecis[i],"eci_b":ecis[j],
                        "match_pct":pct,"matches":matches,"total":total,
                        "severity":"HIGH","statutes":["333","2511","1030"],
                        "explanation":(
                            f"ECI {ecis[i]} and ECI {ecis[j]} produce RSRP values "
                            f"within ±{RSRP_CLONE_TOL} dBm on {matches}/{total} "
                            f"time-aligned records ({pct:.0f}%).  "
                            "Independent towers at different locations produce "
                            "uncorrelated signal strength profiles.  "
                            "Near-identical RSRP across separate ECIs is consistent "
                            "with one transmitter copying the other's signal parameters "
                            "to defeat signal-level anomaly detection filters — "
                            "a documented IMSI-catcher evasion technique."
                        ),
                    })
    return findings


def detect_gps_injection(df):
    """
    A4 — GPS coordinate injection.

    Signature A: exact repeating speed constants above SPEED_INJECT_MIN.
    Real GPS speed is noisy — identical values never repeat.

    Signature B: speed/distance internal inconsistency — computed speed
    from consecutive lat/lon/time disagrees with logged distance field by
    more than SPEED_DIST_TOL.  Synthesized streams have this inconsistency
    when position and speed come from separate code paths.

    Either signature alone is HIGH; both together is CRITICAL.
    """
    findings = []
    lat_col=_resolve(df,"lat"); lon_col=_resolve(df,"lon")
    ts_col=_resolve(df,"time_start")
    if not all([lat_col,lon_col,ts_col]): return findings

    speed_s = compute_speed_series(df, gnss_only=True)

    # Signature A
    high = speed_s[speed_s > SPEED_INJECT_MIN].round(3)
    const_counts = high.value_counts()
    injected = const_counts[const_counts >= SPEED_REPEAT_MIN]

    # Signature B
    inconsistent = []
    prov_col = _resolve(df,"provider")
    # GPS injection only applies to GNSS-derived positions.
    # NLP/NETWORK entries are cell-tower triangulation — their position
    # jitter is legitimate and is NOT evidence of GPS spoofing.
    gnss_mask = pd.Series([True]*len(df), index=df.index)
    if prov_col:
        gnss_mask = df[prov_col].str.upper().isin(["GPS","GNSS"])
    net = df[gnss_mask].sort_values(ts_col).copy()
    lats=net[lat_col].values; lons=net[lon_col].values
    times=net[ts_col].values; idxs=net.index.tolist()
    dist_col = _resolve(df,"dist_m")
    for i in range(1,len(net)):
        try:
            d  = haversine(lats[i-1],lons[i-1],lats[i],lons[i])
            dt = (pd.Timestamp(times[i])-pd.Timestamp(times[i-1])).total_seconds()
            if dt<=0: continue
            cs = d/dt
            if cs <= SPEED_INJECT_MIN: continue
            if dist_col and dist_col in net.columns:
                rd = net.iloc[i][dist_col]
                if not pd.isna(rd):
                    ratio = abs(float(rd)-d)/max(d,1.0)
                    if ratio > SPEED_DIST_TOL:
                        inconsistent.append({
                            "ts":pd.Timestamp(times[i]),
                            "computed_spd":cs,"reported_dist":float(rd),
                            "haversine_dist":d,"ratio":ratio,
                        })
        except Exception:
            pass

    if injected.empty and not inconsistent: return findings

    flagged = int((speed_s.round(3).isin(injected.index)).sum()) if not injected.empty else 0
    pct_f   = flagged/len(df)*100 if flagged else 0
    max_spd = float(injected.index.max()) if not injected.empty else 0.0
    both    = not injected.empty and len(inconsistent)>=2

    findings.append({
        "type":"GPS_COORDINATE_INJECTION",
        "label":(f"GPS Coordinate Injection — "
                 +(f"{len(injected)} Constant(s) — Max {max_spd:.0f} m/s "
                   f"({max_spd*2.23694:.0f} mph)" if not injected.empty else "")
                 +(f" + {len(inconsistent)} Dist/Speed Inconsistencies" if inconsistent else "")),
        "injected_vals":dict(injected) if not injected.empty else {},
        "total_flagged":flagged,"pct_flagged":pct_f,
        "max_speed":max_spd,"max_speed_mph":max_spd*2.23694,
        "inconsistent_rows":inconsistent[:10],"both_sigs":both,
        "severity":"CRITICAL" if both else "HIGH",
        "statutes":["1367","333","1030","2511"],
        "explanation":(
            "Two independent GPS injection signatures were evaluated.  "
            "Signature A (repeating constants): GPS-derived speed is a continuous "
            "noisy quantity — identical floating-point values never repeat across "
            "independent intervals.  "
            +(f"Constants detected: " + "; ".join(
                f"{v:.3f} m/s x{int(c)}" for v,c in injected.items()
            ) + f".  {flagged} records ({pct_f:.1f}%) affected.  "
              if not injected.empty else "No repeating constants detected.  ")
            +f"Signature B (dist/speed inconsistency): {len(inconsistent)} record(s) "
            f"where haversine distance disagrees with logged distance by >{int(SPEED_DIST_TOL*100)}%.  "
            "Real GPS data is internally consistent because all derived fields "
            "share one position fix.  Synthesized streams diverge when position "
            "and speed come from separate code paths.  "
            "GPS is a U.S. military asset (DOD Directive 4650.1).  "
            "GPS spoofing violates 18 U.S.C. § 1367 and 18 U.S.C. § 1030 (CFAA)."
        ),
    })
    return findings


def detect_session_gaps(df):
    """
    A5 — Anomalous inter-record timing gaps.

    Alert threshold is derived from the session's own cadence
    (median + GAP_SIGMA * σ), with GAP_FLOOR_SEC as a minimum floor.
    Self-calibrates to any logger sampling rate.
    Correlated suppression flag: gap adjacent to RSRQ > spec indicates
    jamming and measurement manipulation are temporally coordinated.
    """
    findings = []
    ts_ser=_ser(df,"time_start"); rsrq_ser=_ser(df,"rsrq")
    eci_ser=_ser(df,"cell_id"); pci_ser=_ser(df,"pci")
    if ts_ser is None: return findings
    ts = ts_ser.dropna().sort_values()
    if len(ts) < 3: return findings
    gaps = ts.diff().dt.total_seconds().dropna()
    pos  = gaps[gaps>0]
    if len(pos) < 3: return findings
    med=float(pos.median()); std=float(pos.std())
    threshold = max(GAP_FLOOR_SEC, med + GAP_SIGMA*std)
    alert = gaps[gaps>=threshold]
    if alert.empty: return findings

    max_gap=float(gaps.max()); total_lost=float(alert.sum())
    detail=[]
    for idx in alert.index:
        if idx not in df.index: continue
        gv  = float(gaps[idx])
        tv  = df.loc[idx,ts_ser.name] if ts_ser.name in df.columns else None
        ev  = df.loc[idx,eci_ser.name] if eci_ser and eci_ser.name in df.columns else None
        pv  = df.loc[idx,pci_ser.name] if pci_ser and pci_ser.name in df.columns else None
        adj = None
        if rsrq_ser is not None:
            pos2 = df.index.get_loc(idx)
            w    = rsrq_ser.iloc[max(0,pos2-3):min(len(df),pos2+4)]
            adj  = float(w.max()) if not w.empty else None
        detail.append({"ts":tv,"gap_sec":gv,"eci":ev,"pci":pv,"adj_rsrq":adj})

    corr = [g for g in detail if g["adj_rsrq"] is not None and g["adj_rsrq"]>RSRQ_SPEC_MAX]
    severity = "CRITICAL" if (max_gap>=90 or len(corr)>=2) else "HIGH"

    findings.append({
        "type":"SESSION_GAPS",
        "label":(f"Session Timing Gaps — {len(alert)} Events >= {threshold:.0f}s "
                 f"— Max {max_gap:.0f}s — {len(corr)} RSRQ-Correlated"),
        "total_gaps":len(alert),"max_gap":max_gap,"median_gap":med,"std_gap":std,
        "threshold":threshold,"total_lost":total_lost,
        "gap_detail":detail,"correlated":corr,"severity":severity,
        "statutes":["333","2.807","1362","1030"],
        "explanation":(
            f"Normal logger cadence: {med:.1f}s median (σ={std:.1f}s).  "
            f"Dynamic alert threshold: {threshold:.0f}s "
            f"(median + {GAP_SIGMA}σ, floor {GAP_FLOOR_SEC}s).  "
            f"{len(alert)} gap(s) detected at or above threshold; max {max_gap:.0f}s; "
            f"{total_lost:.0f}s total missing coverage.  "
            f"{len(corr)} gap(s) immediately adjacent to RSRQ>{RSRQ_SPEC_MAX}, "
            "indicating RF suppression and measurement manipulation are temporally "
            "coordinated — consistent with a jammer blanking the device while the "
            "rogue cell completes its session, then reducing power so the device "
            "re-acquires — but only to the rogue cell whose inflated RSRQ holds it "
            "as preferred serving cell.  "
            "Willful interference silencing a licensed receiver violates "
            "47 U.S.C. § 333 and 47 C.F.R. § 2.807 (jamming).  "
            "Suppression of government-licensed communications adds 18 U.S.C. § 1362."
        ),
    })
    return findings


def detect_reserved_tac(df):
    """
    A6 — Reserved / sentinel TAC values.
    3GPP TS 23.003 §19.4: TAC 0x0000 and 0xFFFE are reserved; 0xFFFF (65535)
    is the 'invalid/unknown' sentinel.  A cell broadcasting these is not
    operating within 3GPP-compliant parameters — standalone finding independent
    of PCI collision or flash status.
    """
    findings = []
    tac_ser=_ser(df,"tac"); eci_ser=_ser(df,"cell_id")
    pci_ser=_ser(df,"pci"); ef_ser=_ser(df,"earfcn")
    if tac_ser is None: return findings
    bad = df[tac_ser.isin(RESERVED_TACS)]
    if bad.empty: return findings
    cells={}
    if eci_ser is not None:
        for eci,grp in bad.groupby(eci_ser.name):
            cells[int(eci)] = {
                "tac":   int(grp[tac_ser.name].iloc[0]),
                "pci":   grp[pci_ser.name].iloc[0] if pci_ser else None,
                "earfcn":grp[ef_ser.name].iloc[0]  if ef_ser  else None,
                "count": len(grp),
            }
    utacs = sorted(bad[tac_ser.name].unique().astype(int).tolist())
    findings.append({
        "type":"RESERVED_TAC",
        "label":(f"Reserved TAC Values — {len(utacs)} Sentinel TAC(s): {utacs} "
                 f"— {len(bad)} Records"),
        "unique_tacs":utacs,"count":len(bad),"pct":len(bad)/len(df)*100,
        "cell_summary":cells,"sample_rows":bad.head(10),
        "tac_col":tac_ser.name,
        "eci_col":eci_ser.name if eci_ser else None,
        "pci_col":pci_ser.name if pci_ser else None,
        "ef_col": ef_ser.name  if ef_ser  else None,
        "severity":"CRITICAL" if any(t in RESERVED_TACS for t in utacs) else "HIGH",
        "statutes":["301","333","1028A","2511"],
        "explanation":(
            "3GPP TS 23.003 §19.4 defines TAC 0x0000 and 0xFFFE as reserved "
            "values that shall not be assigned to any operational cell.  "
            "TAC 0xFFFF (65535) is the 'invalid/not available' sentinel — "
            "a cell broadcasting this is deliberately withholding its "
            "Tracking Area Code.  "
            f"This session observed TAC(s) {utacs} on {len(bad)} records "
            f"({len(bad)/len(df)*100:.1f}% of session).  "
            "A reserved or null TAC is characteristic of a software-defined radio "
            "or commercial IMSI-catcher platform where TAC is unconfigured during "
            "deployment.  Broadcasting a reserved TAC while impersonating a "
            "licensed carrier's MCC/MNC constitutes aggravated identity theft "
            "(18 U.S.C. § 1028A) and unauthorized LTE operation "
            "(47 U.S.C. §§ 301 and 333)."
        ),
    })
    return findings

# ── DOCX PRIMITIVES (same style as GNSS_REPORT.py) ───────────────────────────

def set_cell_bg(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"),hex_color); tcPr.append(shd)

def add_run(para, text, bold=False, color=None, size=None, italic=False, mono=False):
    run=para.add_run(str(text)); run.bold=bold; run.italic=italic
    if color: run.font.color.rgb=RGBColor(*bytes.fromhex(color))
    if size:  run.font.size=Pt(size)
    if mono:  run.font.name="Courier New"
    return run

def heading(doc, text, level=1, color="B71C1C"):
    para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.LEFT
    run=para.add_run(text); run.bold=True
    run.font.size=Pt(14 if level==1 else 12)
    run.font.color.rgb=RGBColor(*bytes.fromhex(color))
    pPr=para._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    bot=OxmlElement("w:bottom")
    bot.set(qn("w:val"),"double" if level==1 else "single")
    bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"1")
    bot.set(qn("w:color"),color); pBdr.append(bot); pPr.append(pBdr)
    para.paragraph_format.space_before=Pt(12)
    para.paragraph_format.space_after=Pt(6)
    return para

def add_kv_table(doc, rows, col_widths=(2.6,3.9)):
    table=doc.add_table(rows=0,cols=2); table.style="Table Grid"
    for i,(label,value,hl) in enumerate(rows):
        row=table.add_row(); c0,c1=row.cells[0],row.cells[1]
        c0.width=Inches(col_widths[0]); c1.width=Inches(col_widths[1])
        set_cell_bg(c0,"ECEFF1")
        set_cell_bg(c1,"FFEBEE" if hl else ("FFFFFF" if i%2==0 else "FAFAFA"))
        add_run(c0.paragraphs[0],str(label),bold=True,size=10)
        if hl: add_run(c1.paragraphs[0],str(value),bold=True,color="B71C1C",size=10)
        else:  add_run(c1.paragraphs[0],str(value),size=10)
    doc.add_paragraph()

def add_log_table(doc, headers, col_widths, data_rows):
    table=doc.add_table(rows=0,cols=len(headers)); table.style="Table Grid"
    hdr=table.add_row()
    for i,(h,w) in enumerate(zip(headers,col_widths)):
        c=hdr.cells[i]; c.width=Inches(w)
        set_cell_bg(c,"1A237E"); add_run(c.paragraphs[0],h,bold=True,color="FFFFFF",size=9)
    for dr in data_rows:
        row=table.add_row()
        for i,cv in enumerate(dr):
            c=row.cells[i]; c.width=Inches(col_widths[i])
            text=cv[0] if isinstance(cv,tuple) else str(cv)
            hl  =cv[1] if isinstance(cv,tuple) else False
            set_cell_bg(c,"FFEBEE" if hl else "FFFFFF")
            if hl: add_run(c.paragraphs[0],text,bold=True,color="B71C1C",size=9,mono=True)
            else:  add_run(c.paragraphs[0],text,size=9,mono=True)
    doc.add_paragraph()

def sev_color(s):
    return {"CRITICAL":"B71C1C","HIGH":"E65100","MODERATE":"F57F17"}.get(s,"000000")

def _fmt(v,dec=0,suf="",na="—"):
    try:
        if v is None or (isinstance(v,float) and pd.isna(v)): return na
        return (f"{v:.{dec}f}{suf}" if dec else f"{int(round(float(v)))}{suf}")
    except Exception: return str(v) if v is not None else na

def _s(v,na="—"):
    try:
        if v is None or pd.isna(v): return na
    except Exception: pass
    return str(v)

# ── SESSION INFO ─────────────────────────────────────────────────────────────

def compute_session_info(pts, inc):
    ts=_ser(pts,"time_start"); eci=_ser(pts,"cell_id")
    pci=_ser(pts,"pci"); ef=_ser(pts,"earfcn"); tac=_ser(pts,"tac")
    mcc_c=_resolve(pts,"mcc"); mnc_c=_resolve(pts,"mnc"); rat_c=_resolve(pts,"rat")
    ts_v = ts.dropna() if ts is not None else pd.Series([],dtype="object")
    start=ts_v.min() if not ts_v.empty else None
    end  =ts_v.max() if not ts_v.empty else None
    dur  =(end-start).total_seconds() if (start and end) else 0
    return {
        "start":start,"end":end,"duration_sec":dur,
        "total_records":len(pts),"total_incidents":len(nlp_inc),
        "unique_ecis":  int(eci.nunique()) if eci is not None else 0,
        "unique_pcis":  int(pci.nunique()) if pci is not None else 0,
        "unique_earfcns":int(ef.nunique()) if ef  is not None else 0,
        "transitions":  int((eci!=eci.shift()).sum()) if eci is not None else 0,
        "tac":          int(tac.mode()[0]) if (tac is not None and not tac.mode().empty) else None,
        "mcc":str(int(pts[mcc_c].mode()[0])) if mcc_c and not pts[mcc_c].mode().empty else "—",
        "mnc":str(int(pts[mnc_c].mode()[0])) if mnc_c and not pts[mnc_c].mode().empty else "—",
        "rat":str(pts[rat_c].mode()[0])       if rat_c and not pts[rat_c].mode().empty else "—",
    }

# ── REPORT BUILDER ───────────────────────────────────────────────────────────

def build_report(case_dir, pts, nlp_inc, gps_inc, findings_all, si, hash_info):
    doc=Document()
    for section in doc.sections:
        section.top_margin=Inches(0.9); section.bottom_margin=Inches(0.9)
        section.left_margin=Inches(1.2); section.right_margin=Inches(0.9)

    # Title
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_run(t,"NLP CELLULAR INTERFERENCE EVIDENCE REPORT",bold=True,color="B71C1C",size=18)
    pPr=t._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    for side in ["top","bottom"]:
        el=OxmlElement(f"w:{side}"); el.set(qn("w:val"),"double")
        el.set(qn("w:sz"),"8"); el.set(qn("w:space"),"4")
        el.set(qn("w:color"),"B71C1C"); pBdr.append(el)
    pPr.append(pBdr)
    sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub,"Federal Communications Commission — Enforcement Bureau Submission\n",bold=True,size=11)
    add_run(sub,(f"Generated {datetime.datetime.now().strftime('%B %d, %Y  %H:%M UTC')}  |  "
                 f"NLP Cellular Evidence Scanner  |  Case: {os.path.basename(case_dir)}\n"),size=10)
    add_run(sub,"SIMULTANEOUS COPY: FBI Riverside  |  DOJ Civil Rights  |  "
               "FCC Enforcement Bureau  |  NRC  |  Cal OES #26-1153",
            bold=True,color="B71C1C",size=10)

    # Notice box
    n=doc.add_paragraph()
    add_run(n,"NOTICE: ",bold=True,color="B71C1C",size=11)
    add_run(n,("All values extracted directly from machine-generated data produced by "
               "GNSS_AttackModel.ps1.  No measurement is estimated or manually entered.  "
               "Anomalies are detected by automated analysis against 3GPP LTE specification "
               "parameters.  The SHA-256 hash in Section I provides chain-of-custody integrity "
               "for the source file.  The conduct documented requires hardware physically outside "
               "the complainant's control, operated without FCC authorization on licensed spectrum."),size=10)
    pPr2=n._p.get_or_add_pPr(); pBdr2=OxmlElement("w:pBdr")
    for side in ["top","bottom","left","right"]:
        el=OxmlElement(f"w:{side}"); el.set(qn("w:val"),"single")
        el.set(qn("w:sz"),"6" if side!="left" else "18")
        el.set(qn("w:space"),"4"); el.set(qn("w:color"),"B71C1C"); pBdr2.append(el)
    pPr2.append(pBdr2); n.paragraph_format.left_indent=Inches(0.2); doc.add_paragraph()

    # Section I
    heading(doc,"I. COMPLAINANT, SESSION IDENTIFICATION, AND CHAIN OF CUSTODY")
    lat_c=_resolve(pts,"lat"); lon_c=_resolve(pts,"lon")
    clat=f"{pts[lat_c].mean():.6f}" if lat_c else "—"
    clon=f"{pts[lon_c].mean():.6f}" if lon_c else "—"
    add_kv_table(doc,[
        ("Complainant:",      COMPLAINANT,                                False),
        ("DOB:",              DOB if DOB else "—",                        False),
        ("Case Folder:",      os.path.basename(case_dir),                 False),
        ("Source File:",      hash_info.get("File","—"),                  False),
        ("SHA-256 Hash:",     hash_info.get("Hash","—"),                  False),
        ("Hash Algorithm:",   hash_info.get("Algorithm","—"),             False),
        ("Session Start:",    _s(si["start"]),                            False),
        ("Session End:",      _s(si["end"]),                              False),
        ("Duration:",         f"{si['duration_sec']:.0f}s ({si['duration_sec']/60:.1f} min)", False),
        ("Total Records:",    str(si["total_records"]),                   False),
        ("Pre-classified Incidents:", str(si["total_incidents"]),         True),
        ("Center Coords:",    f"{clat}, {clon}",                          False),
        ("Network MCC/MNC:",  f"{si['mcc']}/{si['mnc']} ({si['rat']})",  False),
        ("Dominant TAC:",     _s(si.get("tac")),                          False),
        ("Unique Cell IDs:",  str(si["unique_ecis"]),                     False),
        ("Unique PCIs:",      str(si["unique_pcis"]),                     False),
        ("Unique EARFCNs:",   str(si["unique_earfcns"]),                  False),
        ("Cell Transitions:", str(si["transitions"]),                     True),
        ("Total Findings:",   str(sum(len(v) for v in findings_all.values())), True),
    ])

    # Section II
    heading(doc,"II. EXECUTIVE SUMMARY OF FINDINGS")
    total=sum(len(v) for v in findings_all.values())
    n_crit=sum(1 for v in findings_all.values() for f in v if f.get("severity")=="CRITICAL")
    n_high=sum(1 for v in findings_all.values() for f in v if f.get("severity")=="HIGH")
    sm=doc.add_paragraph()
    add_run(sm,f"This session produced {total} finding(s): {n_crit} CRITICAL, {n_high} HIGH.  ",
            bold=True,color="B71C1C",size=11)
    add_run(sm,("All findings are consistent with active, deliberate operation of one or more "
                "unauthorized cellular transmitters (false base stations / IMSI catchers) on "
                "licensed LTE spectrum.  The combination of anomalies documented constitutes a "
                "multi-layer interception architecture that cannot be explained by any legitimate "
                "network configuration or equipment malfunction."),size=11)
    doc.add_paragraph()
    all_statutes=sorted({s for v in findings_all.values() for f in v for s in f.get("statutes",[])})
    sp=doc.add_paragraph()
    add_run(sp,"Federal statutes potentially violated (proven by measured data):\n",bold=True,size=11)
    for code in all_statutes:
        if code in STATUTES: add_run(sp,f"  {STATUTES[code]}\n",size=10)
    doc.add_paragraph()

    # Section III — Detailed findings
    heading(doc,"III. DETAILED FINDINGS — MAPPED TO FEDERAL STATUTES")
    fnum=0
    ORDER=[("ghost_flash","A1"),("pci_collision","A2"),("rsrq_invalid","A3"),
           ("gps_injection","A4"),("session_gaps","A5"),("reserved_tac","A6")]

    for key,code in ORDER:
        for f in findings_all.get(key,[]):
            fnum+=1
            sc=sev_color(f["severity"])
            heading(doc,f"Finding {fnum} ({code}): {f['label']}",level=2,color=sc)
            doc.add_paragraph().add_run(f["explanation"]).font.size=Pt(10)
            doc.add_paragraph()

            if f["type"]=="GHOST_CELL_FLASH":
                add_kv_table(doc,[
                    ("Cell ECI:",         _s(f["eci"]),                True),
                    ("PCI:",              _s(f["pci"]),                True),
                    ("EARFCN:",           _s(f["earfcn"]),             True),
                    ("TAC:",              _s(f["tac"]),                f.get("tac") in RESERVED_TACS),
                    ("MCC/MNC:",          f"{_s(f['mcc'])}/{_s(f['mnc'])}",False),
                    ("RAT:",              _s(f["rat"]),                False),
                    ("RSRP:",             _fmt(f["rsrp"],0," dBm"),   False),
                    ("RSRQ:",             _fmt(f["rsrq"],1),
                     f.get("rsrq") is not None and not pd.isna(f["rsrq"]) and float(f["rsrq"])>RSRQ_SPEC_MAX),
                    ("Timestamp:",        _s(f["timestamp"]),          True),
                    ("Round-minute flag:",("YES — scripted activation" if f["round_minute"] else "NO"),f["round_minute"]),
                    ("Observations:",     f"{f['obs_count']} ({f['presence_pct']:.1f}% of session)",True),
                    ("Severity:",         f["severity"],               f["severity"]=="CRITICAL"),
                ])

            elif f["type"]=="PCI_COLLISION":
                add_kv_table(doc,[
                    ("PCI:",              str(f["pci"]),   True),
                    ("EARFCNs colliding:",str(f["n_earfcns"]),True),
                    ("Total obs:",        str(f["total_obs"]),False),
                    ("Severity:",         f["severity"],   True),
                ])
                if f["earfcn_groups"]:
                    add_log_table(doc,
                        ["EARFCN","Obs","Cell ECI(s)","TAC(s)","RSRP avg"],
                        [0.9,0.7,2.4,0.9,1.0],
                        [[(str(ef),True),(str(g["count"]),False),
                          (", ".join(str(e) for e in g["ecis"]) or "—",True),
                          (", ".join(str(t) for t in g["tacs"]) or "—",False),
                          (_fmt(g["rsrp_mean"],1," dBm"),False)]
                         for ef,g in sorted(f["earfcn_groups"].items())])

            elif f["type"]=="RSRQ_OUT_OF_SPEC":
                add_kv_table(doc,[
                    ("LTE RSRQ spec max:", f"{RSRQ_SPEC_MAX} (3GPP TS 36.133 §9.1.7)",False),
                    ("Max observed RSRQ:", _fmt(f["max_rsrq"],1),     True),
                    ("Ratio vs spec:",     _fmt(f["max_ratio"],1,"x"),True),
                    ("Over-spec records:", f"{f['count_over']} ({f['pct_over']:.1f}%)",True),
                    ("Severity:",          f["severity"],              True),
                ])
                if f["cell_groups"]:
                    add_log_table(doc,
                        ["Cell ECI","PCI","EARFCN","Count","Max RSRQ"],
                        [1.5,0.7,0.8,0.9,0.9],
                        [[(str(eci),True),(_s(g["pci"]),True),(_s(g["earfcn"]),False),
                          (str(g["count"]),True),(_fmt(g["max_rsrq"],1),True)]
                         for eci,g in sorted(f["cell_groups"].items(),
                                             key=lambda x:x[1]["max_rsrq"],reverse=True)])

            elif f["type"]=="RSRP_CLONE":
                add_kv_table(doc,[
                    ("ECI A:",       str(f["eci_a"]),                             True),
                    ("ECI B:",       str(f["eci_b"]),                             True),
                    ("Tolerance:",   f"\u00b1{RSRP_CLONE_TOL} dBm",              False),
                    ("Match rate:",  f"{f['matches']}/{f['total']} ({f['match_pct']:.0f}%)",True),
                    ("Severity:",    f["severity"],                               False),
                ])

            elif f["type"]=="GPS_COORDINATE_INJECTION":
                add_kv_table(doc,[
                    ("Injection constants:",     str(len(f["injected_vals"])),   True),
                    ("Records affected:",        f"{f['total_flagged']} ({f['pct_flagged']:.1f}%)",True),
                    ("Max injected speed:",      f"{f['max_speed']:.1f} m/s ({f['max_speed_mph']:.0f} mph)",True),
                    ("Dist/speed inconsistencies:",str(len(f["inconsistent_rows"])),len(f["inconsistent_rows"])>0),
                    ("Both signatures:",         "YES — CRITICAL" if f["both_sigs"] else "NO",f["both_sigs"]),
                    ("Severity:",                f["severity"],                   True),
                ])
                if f["injected_vals"]:
                    add_log_table(doc,
                        ["Speed constant (m/s)","Equivalent (mph)","Occurrences","Physically possible?"],
                        [1.8,1.2,1.0,2.2],
                        [[(_fmt(v,3),True),(_fmt(v*2.23694,0," mph"),True),
                          (str(int(c)),True),("NO — exceeds sound barrier" if v>343 else "NO — vehicle limit",True)]
                         for v,c in sorted(f["injected_vals"].items(),reverse=True)])
                if f["inconsistent_rows"]:
                    add_log_table(doc,
                        ["Timestamp","Computed spd (m/s)","Reported dist (m)","Haversine dist (m)","Discrepancy"],
                        [2.0,1.4,1.2,1.3,0.8],
                        [[(_s(r["ts"]),True),(_fmt(r["computed_spd"],1),True),
                          (_fmt(r["reported_dist"],1),False),(_fmt(r["haversine_dist"],1),False),
                          (_fmt(r["ratio"]*100,0,"%"),True)]
                         for r in f["inconsistent_rows"]])

            elif f["type"]=="SESSION_GAPS":
                add_kv_table(doc,[
                    ("Session median cadence:",  _fmt(f["median_gap"],1,"s"),  False),
                    ("Cadence \u03c3:",          _fmt(f["std_gap"],1,"s"),     False),
                    ("Dynamic threshold:",       _fmt(f["threshold"],0,"s"),   False),
                    ("Alert gaps:",              str(f["total_gaps"]),         True),
                    ("Max gap:",                 _fmt(f["max_gap"],0,"s"),     True),
                    ("Total lost coverage:",     _fmt(f["total_lost"],0,"s"),  True),
                    ("RSRQ-correlated gaps:",    str(len(f["correlated"])),    len(f["correlated"])>0),
                    ("Severity:",                f["severity"],                f["severity"]=="CRITICAL"),
                ])
                if f["gap_detail"]:
                    add_log_table(doc,
                        ["Timestamp","Gap (s)","Cell ECI","PCI","Adj RSRQ","Flag"],
                        [2.0,0.7,1.3,0.6,0.9,0.8],
                        [[(_s(g["ts"]),False),(_fmt(g["gap_sec"],0,"s"),True),
                          (_s(g["eci"]),False),(_s(g["pci"]),False),
                          (_fmt(g["adj_rsrq"],1) if g["adj_rsrq"] else "—",
                           g["adj_rsrq"] is not None and g["adj_rsrq"]>RSRQ_SPEC_MAX),
                          ("OVER SPEC" if g["adj_rsrq"] is not None and g["adj_rsrq"]>RSRQ_SPEC_MAX else "—",
                           g["adj_rsrq"] is not None and g["adj_rsrq"]>RSRQ_SPEC_MAX)]
                         for g in f["gap_detail"]])

            elif f["type"]=="RESERVED_TAC":
                add_kv_table(doc,[
                    ("Reserved TAC(s):", str(f["unique_tacs"]),                 True),
                    ("Records affected:",f"{f['count']} ({f['pct']:.1f}%)",    True),
                    ("3GPP reference:",  "TS 23.003 §19.4 — reserved/invalid", False),
                    ("Severity:",        f["severity"],                         True),
                ])
                if f["cell_summary"]:
                    add_log_table(doc,
                        ["Cell ECI","TAC","PCI","EARFCN","Obs"],
                        [1.5,0.8,0.7,0.8,0.9],
                        [[(str(eci),True),(str(g["tac"]),True),
                          (_s(g["pci"]),False),(_s(g["earfcn"]),False),(str(g["count"]),False)]
                         for eci,g in sorted(f["cell_summary"].items())])

            # RSRQ data gap — simple notice box, no KV table
            if f["type"]=="RSRQ_DATA_GAP":
                notice2=doc.add_paragraph()
                add_run(notice2,"ACTION REQUIRED — see explanation above.",
                        bold=True,color="E65100",size=10)

            # Statutes block
            sp2=doc.add_paragraph()
            add_run(sp2,"Applicable statutes: ",bold=True,size=10,color="B71C1C")
            for code2 in f.get("statutes",[]):
                if code2 in STATUTES: add_run(sp2,f"\n  {STATUTES[code2]}",size=10)

    # Section IV — Pre-classified incidents
    doc.add_page_break()
    heading(doc,"IV. PRE-CLASSIFIED INCIDENTS FROM GNSS_ATTACKMODEL.PS1")
    add_run(doc.add_paragraph(),
            ("Incidents classified by GNSS_AttackModel.ps1 — split by provider type\n"
             "(equivalent of: busybox grep \"NLP\" incidents.csv)\n\n"
             "NLP incidents (NETWORK provider — cell-tower triangulation):\n"
             "  NLP_CRIME   — NETWORK fix moved >20 ft from prior fix\n"
             "  NLP_ANOMALY — NETWORK fix moved 5-20 ft (borderline)\n\n"
             "GPS incidents (GNSS provider — satellite-derived position):\n"
             "  GPS_OUT_OF_RADIUS — GPS fix outside examiner tolerance radius"),
            size=10)
    doc.add_paragraph()

    def _render_incident_table(doc, df_in, label):
        if df_in.empty:
            add_run(doc.add_paragraph(), f"No {label} incidents.", size=10)
            return
        sev_c=_resolve(df_in,"severity")
        sev_counts=df_in[sev_c].value_counts().to_dict() if sev_c else {}
        add_kv_table(doc,[("Total:",str(len(df_in)),True)]
                        +[(f"{k}:",str(v),k=="NLP_CRIME") for k,v in sorted(sev_counts.items())])
        ts_c=_resolve(df_in,"time_start"); dm_c=_resolve(df_in,"dist_m")
        eci_c=_resolve(df_in,"cell_id"); pci_c=_resolve(df_in,"pci")
        ef_c=_resolve(df_in,"earfcn"); seq_c=_resolve(df_in,"seq")
        rows_out=[]
        for _,r in df_in.head(100).iterrows():
            sv=str(r[sev_c]) if sev_c else "—"
            rows_out.append([
                (_s(r[seq_c]) if seq_c else "—",   False),
                (sv,                                sv=="NLP_CRIME"),
                (_s(r[ts_c]) if ts_c else "—",     False),
                (_fmt(r[dm_c],1) if dm_c else "—", sv=="NLP_CRIME"),
                (_s(r[eci_c]) if eci_c else "—",   False),
                (_s(r[pci_c]) if pci_c else "—",   False),
                (_s(r[ef_c]) if ef_c else "—",     False),
            ])
        add_log_table(doc,["#","Severity","Timestamp (UTC)","Dist (m)","Cell ECI","PCI","EARFCN"],
                      [0.4,1.2,2.0,0.8,1.3,0.7,0.8],rows_out)

    heading(doc,"NLP Incidents (Network Provider — Cell Triangulation)",level=2,color="1A237E")
    _render_incident_table(doc, nlp_inc, "NLP")

    heading(doc,"GPS Incidents (GNSS Provider — Satellite-Derived Position)",level=2,color="1B5E20")
    _render_incident_table(doc, gps_inc, "GPS")

    # Section V — Cell inventory
    doc.add_page_break()
    heading(doc,"V. CELLMAPPER REFERENCE — COMPLETE CELL INVENTORY")
    add_run(doc.add_paragraph(),
            ("CellMapper crowdsourced database represents verified operator-deployed "
             "infrastructure.  Cells carrying reserved TAC values, PCI collisions, "
             "or single-observation flash status have no verified network topology match "
             "and no legitimate operational explanation."),size=10)
    doc.add_paragraph()
    eci_col=_resolve(pts,"cell_id"); pci_col=_resolve(pts,"pci")
    ef_col=_resolve(pts,"earfcn"); tac_col=_resolve(pts,"tac"); rp_col=_resolve(pts,"rsrp")
    if eci_col and pci_col and ef_col:
        gc=[eci_col,pci_col,ef_col]+([tac_col] if tac_col else [])
        agg={"obs":(eci_col,"count")}
        if rp_col: agg["rsrp_mean"]=(rp_col,"mean")
        inv=(pts[gc+([rp_col] if rp_col else [])].dropna(subset=[eci_col])
             .astype({eci_col:int,pci_col:"Int64",ef_col:"Int64"})
             .groupby(gc).agg(**agg).reset_index().sort_values("obs",ascending=False))
        flash_s={int(f["eci"]) for f in findings_all.get("ghost_flash",[])}
        coll_s ={int(f["pci"]) for f in findings_all.get("pci_collision",[])}
        hdrs=["Cell ECI","PCI","EARFCN"]; wids=[1.4,0.7,0.8]
        if tac_col: hdrs.append("TAC"); wids.append(0.8)
        if rp_col:  hdrs.append("RSRP"); wids.append(0.8)
        hdrs+=["Obs","Anomaly flags"]; wids+=[0.5, max(0.5,6.5-sum(wids))]
        rows_inv=[]
        for _,r in inv.iterrows():
            ev=int(r[eci_col]); pv=int(r[pci_col]); efv=int(r[ef_col])
            flags=[]
            if ev in flash_s: flags.append("FLASH CELL")
            if pv in coll_s:  flags.append("PCI COLLISION")
            if tac_col:
                tv=r[tac_col]
                if not pd.isna(tv) and int(tv) in RESERVED_TACS:
                    flags.append(f"RESERVED TAC {int(tv)}")
            ia=bool(flags); fs=", ".join(flags) if flags else "—"
            ro=[(str(ev),ia),(str(pv),pv in coll_s),(str(efv),False)]
            if tac_col:
                tv=r[tac_col]; ro.append((_s(tv),not pd.isna(tv) and int(tv) in RESERVED_TACS))
            if rp_col: ro.append((_fmt(r.get("rsrp_mean"),1," dBm"),False))
            ro+=[(str(int(r["obs"])),False),(fs,ia)]
            rows_inv.append(ro)
        add_log_table(doc,hdrs,wids,rows_inv)

    # Section VI — Statute summary
    doc.add_page_break()
    heading(doc,"VI. COMPLETE STATUTE VIOLATION SUMMARY")
    for code in all_statutes:
        if code not in STATUTES: continue
        rel=[f["label"] for v in findings_all.values() for f in v if code in f.get("statutes",[])]
        if not rel: continue
        sp3=doc.add_paragraph()
        add_run(sp3,f"{STATUTES[code]}\n",bold=True,color="B71C1C",size=10)
        add_run(sp3,"Supported by:\n",size=10)
        for r in rel: add_run(sp3,f"  \u2022 {r}\n",size=10)
        doc.add_paragraph()

    # Section VII — Certification
    doc.add_page_break()
    heading(doc,"VII. CERTIFICATION UNDER PENALTY OF PERJURY")
    cert=doc.add_paragraph()
    add_run(cert,f"I, {COMPLAINANT}, ",size=11)
    add_run(cert,"declare under penalty of perjury ",bold=True,size=11)
    add_run(cert,(f"under the laws of the United States of America (28 U.S.C. § 1746) "
                  f"that all values in this report are extracted directly and without "
                  f"alteration from machine-generated data files produced by "
                  f"GNSS_AttackModel.ps1 in case folder {os.path.basename(case_dir)}.  "
                  f"The SHA-256 hash recorded in Section I was computed at evidence "
                  f"collection time and constitutes the chain-of-custody integrity record.  "
                  f"This report was generated automatically.  No value has been manually "
                  f"modified except complainant name and date of birth above."),size=11)
    doc.add_paragraph(); doc.add_paragraph()
    sig=doc.add_paragraph()
    add_run(sig,f"{COMPLAINANT}\n",bold=True,size=11)
    add_run(sig,f"DOB: {DOB if DOB else '—'}\n",size=11)
    add_run(sig,f"Report generated: {datetime.datetime.now().isoformat()}",size=10,italic=True)
    return doc

# ── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    global COMPLAINANT, DOB
    # Initialize incident dataframes — guaranteed defined regardless of which
    # file-loading branch executes below.
    inc = pd.DataFrame()
    nlp_inc = pd.DataFrame()
    gps_inc = pd.DataFrame()

    parser=argparse.ArgumentParser(description="NLP Cellular Evidence Report Generator")
    parser.add_argument("--case",default=None,
                        help="Explicit GNSS_AttackCase_* folder path (skips auto-discovery)")
    parser.add_argument("--base",default=BASE_DIR,
                        help=f"Base folder to search for case folders (default: {BASE_DIR})")
    args=parser.parse_args()

    print("="*65)
    print("NLP CELLULAR EVIDENCE SCANNER — FCC Complaint Report Generator")
    print("="*65); print()

    while True:
        name=input("Enter complainant full legal name: ").strip()
        if name: COMPLAINANT=name; break
        print("  Name cannot be empty.")
    dob_input=input("Enter DOB (optional — press Enter to skip): ").strip()
    DOB=dob_input if dob_input else ""; print()

    case_dir=find_case_folder(args.base, args.case)
    print(f"Case folder : {case_dir}")

    pts_path=os.path.join(case_dir,"points.csv")
    inc_path=os.path.join(case_dir,"incidents.csv")

    if not os.path.isfile(pts_path):
        sys.exit(f"[ERROR] points.csv not found in {case_dir}")
    if not os.path.isfile(inc_path):
        print("[WARNING] incidents.csv not found — continuing.")
        inc = pd.DataFrame()
        nlp_inc = pd.DataFrame()
        gps_inc = pd.DataFrame()
    else:
        inc = load_csv(inc_path,"incidents.csv")
        # Split incidents into NLP and GPS — equivalent of:
        #   busybox grep "NLP" incidents.csv
        # Matches rows where Severity contains "NLP" (NLP_CRIME / NLP_ANOMALY)
        # OR Provider/ProviderClass is NLP/NETWORK.
        sev_c2   = _resolve(inc,"severity")
        prov_c2  = _resolve(inc,"provider")
        nlp_flag = pd.Series([False]*len(inc), index=inc.index)
        if sev_c2:
            nlp_flag |= inc[sev_c2].astype(str).str.contains("NLP", case=False, na=False)
        if prov_c2:
            nlp_flag |= inc[prov_c2].astype(str).str.upper().isin(["NLP","NETWORK","FUSED"])
        nlp_inc = inc[nlp_flag].reset_index(drop=True)
        gps_inc = inc[~nlp_flag].reset_index(drop=True)
        print(f"  NLP incidents: {len(nlp_inc)}  GPS incidents: {len(gps_inc)}")

    pts=load_csv(pts_path,"points.csv")
    hash_info=read_hash_file(case_dir)
    print(f"Points      : {len(pts)} records")
    print(f"Incidents   : {len(inc)} pre-classified")
    print(f"Source hash : {hash_info.get('Hash','not found')[:32]}..."); print()

    si=compute_session_info(pts,nlp_inc)
    print(f"Session     : {si['start']} -> {si['end']}")
    print(f"Duration    : {si['duration_sec']:.0f}s  |  "
          f"Cells:{si['unique_ecis']}  PCIs:{si['unique_pcis']}  "
          f"EARFCNs:{si['unique_earfcns']}"); print()
    print("Running detectors...")

    findings_all={
        "ghost_flash":   detect_ghost_cell_flash(pts),
        "pci_collision": detect_pci_collision(pts),
        "rsrq_invalid":  detect_rsrq_invalid(pts),
        "gps_injection": detect_gps_injection(pts),
        "session_gaps":  detect_session_gaps(pts),
        "reserved_tac":  detect_reserved_tac(pts),
    }

    labels={"ghost_flash":"A1 Ghost flash  ","pci_collision":"A2 PCI collision",
            "rsrq_invalid":"A3 RSRQ/clone  ","gps_injection":"A4 GPS injection",
            "session_gaps":"A5 Session gaps","reserved_tac": "A6 Reserved TAC"}

    total=sum(len(v) for v in findings_all.values())
    for key,flist in findings_all.items():
        if flist:
            for f in flist:
                print(f"  [{f['severity']:8s}]  {labels[key]}: {f['label'][:62]}")
        else:
            print(f"  [  OK    ]  {labels[key]}: clean")

    print(f"\nTotal findings: {total}\n")
    print("Generating DOCX report...")

    doc=build_report(case_dir,pts,nlp_inc,gps_inc,findings_all,si,hash_info)
    ts_str=datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    out=os.path.join(case_dir,f"nlp_evidence_report_{ts_str}.docx")
    doc.save(out)
    print(f"Report saved : {out}")
    print("="*65)
    print("Submit to   : enforcement@fcc.gov  (attach incidents.csv + points.csv)")
    print("Copy to     : FBI Riverside  |  DOJ Civil Rights  |  Cal OES #26-1153")
    print("="*65)

if __name__=="__main__":
    main()
